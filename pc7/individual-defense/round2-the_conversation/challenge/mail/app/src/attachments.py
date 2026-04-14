#!/usr/bin/env python3

from faker import Faker
from docx import Document
from io import BytesIO

SEED = 1337

OUT_DIR = "/attachments/"

fake = Faker()
fake.seed_instance(SEED)

def create_docx(contents : list[str]) -> bytes:
    doc = Document()

    for paragraph in contents:
        doc.add_paragraph(paragraph)

    buffer = BytesIO()
    doc.save(buffer)

    return buffer.getvalue()

for i in range(0, 30):
    num_paragraphs = fake.random.randint(3, 5)
    docx = create_docx(fake.paragraphs(nb = num_paragraphs))

    with open(f"{OUT_DIR}{i}.docx", "wb") as f:
        f.write(docx)
