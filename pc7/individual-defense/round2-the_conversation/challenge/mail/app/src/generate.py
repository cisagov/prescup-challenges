#!/usr/bin/env python3

import argparse
import csv

import base64
from docx import Document
from datetime import datetime, timedelta
from email.message import EmailMessage
from email import policy
from io import BytesIO
from faker import Faker
import mailbox

ATTACHMENT_PATH = "/attachments/"

DEFAULT_SEED = 1337

COMPANY_NAME = "Orthanc"
COMPANY_DOMAIN = "orthanc.com"
NUM_NOISE_EMAILS = 5
INTERNAL_EMAIL_TO_PERSON_RATIO = 0.2
EXTERNAL_EMAIL_TO_PERSON_RATIO = 0.8

START_DATE = datetime(2026, 1, 1)
END_DATE = datetime(2026, 1, 31)

VICTIM = {
    "first_name" : "Greg",
    "last_name" : "Collier",
    "email" : f"gcollier@{COMPANY_DOMAIN}"
}

fake = Faker()
employees = [ ]
customers = [ ]

with open("/app/src/internal_messages.csv", "r", encoding="utf-8") as f:
    reader = csv.DictReader(f)
    internal_messages = [{'subject': row['subject'], 'message': row['message']} for row in reader if row.get('subject', '').strip() and row.get('message', '').strip()]

with open("/app/src/external_messages.csv", "r", encoding="utf-8") as f:
    reader = csv.DictReader(f)
    external_messages = [{'subject': row['subject'], 'message': row['message']} for row in reader if row.get('subject', '').strip() and row.get('message', '').strip()]

with open("/payloads/datasheet.docx", "rb") as f:
    payload_attachment_bytes = f.read()

def read_attachment(index) -> bytes:
    with open(f"{ATTACHMENT_PATH}{index}.docx", "rb") as f:
        return f.read()

def roll(probability : float) -> bool:
    return fake.random.random() < probability

def rand_int(min_value : int, max_value : int) -> int:
    return fake.random.randint(min_value, max_value)

def rand_datetime() -> datetime:
    return fake.date_time_between_dates(START_DATE, END_DATE)

def rand_subset(array : list, count : int) -> list:
    return fake.random.sample(array, count)

def select(array : list):
    return fake.random.choice(array)

def select_pair(array : list):
    first = select(array)
    second = select([item for item in array if item != first])
    return first, second

def shuffle(array: list):
    return fake.random.sample(array, len(array))

def rfc2822_date(dt) -> str:
    return dt.strftime("%a, %d %b %Y %H:%M:%S +0000")

def generate_anchor_tag(link : str, text : str | None = None) -> str:
    return f"<a href=\"{link}\">{text if text else link}</a>"

def get_domain_from_email(email : str) -> str:
    return email.split('@')[-1]

def build_auth_header(status : bool, message_id : str, sender_ip : str) -> str:
    if status:
        return (
            f"Authentication-Results: mx.{COMPANY_DOMAIN};"
            f" dkim=pass header.i=@{COMPANY_DOMAIN};"
            f" spf=pass ({COMPANY_DOMAIN}: domain of {message_id} designates {sender_ip} as permitted sender) smtp.mailfrom={message_id};"
            f" dmarc=pass (p=QUARANTINE sp=QUARANTINE dis=NONE) header.from={COMPANY_DOMAIN};"
        )
    else:
        return (
            f"Authentication-Results: mx.{COMPANY_DOMAIN};"
            f" dkim=fail (bad signature) header.i=@{COMPANY_DOMAIN};"
            f" spf=fail ({COMPANY_DOMAIN}: domain of {message_id} does not designate {sender_ip} as permitted sender) smtp.mailfrom={message_id};"
            f" dmarc=fail (p=reject) header.from={COMPANY_DOMAIN};"
        )

def build_dkim_header(sender_domain : str) -> str:
    return (
        f"DKIM-Signature: v=1; a=rsa-sha256; c=relaxed/relaxed;"
        f" d={sender_domain};"
        f" s=dkim;"
        f" h=from:to:subject:message-id;"
        f" bh={base64.b64encode(fake.sha256(raw_output=True)).decode()};"
        f" b={base64.b64encode(fake.binary(length=256)).decode("ascii")};"
    )

def create_msg(sender, recipient, timestamp : datetime, subject : str, body : str, attachment : bytes | None = None, pass_auth : bool = True, message_id : str | None = None) -> EmailMessage:
    sender_domain = get_domain_from_email(sender["email"])
    message_id = message_id if message_id else f"{fake.uuid4()}@{sender_domain}"
    sender_ip = fake.ipv4()

    # msg = EmailMessage(policy=policy.SMTP)
    # msg = EmailMessage(policy=policy.SMTP.clone(refold_source="none", utf8=False))
    msg = EmailMessage(policy=policy.SMTP.clone(refold_source="none", max_line_length=999999, utf8=False))

    # Add all headers in reverse order (as they would be prepended during mail routing)
    extended_headers = [ ]
    
    # Add auth headers
    extended_headers.append(build_auth_header(pass_auth, message_id, sender_ip))
    extended_headers.append(build_dkim_header(sender_domain))

    for h in extended_headers:
        if ":" in h:
            name, value = h.split(":", 1)
            msg[name] = value.lstrip()

    msg["From"] = sender["email"]
    msg["To"] = recipient["email"]
    msg["Date"] = rfc2822_date(timestamp)
    msg["Subject"] = subject
    msg["Message-ID"] = message_id
    msg["Return-Path"] = f"<{sender["email"]}>"

    msg.set_content(body, subtype="html")

    if attachment:
        msg.add_attachment(
            attachment,
            maintype='application',
            subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename='attachment.docx'
        )

    return msg

def create_docx(contents : list[str]) -> bytes:
    doc = Document()

    for paragraph in contents:
        doc.add_paragraph(paragraph)

    buffer = BytesIO()
    doc.save(buffer)

    return buffer.getvalue()



def generate_noise() -> list[EmailMessage]:
    messages = [ ]
    attachment_index = 0

    # customer_service_reps = rand_subset(employees, 9)
    # customer_service_reps.append(VICTIM)

    for external_message in external_messages:
        # Select a random customer
        employee = VICTIM
        customer = select(customers)

        timestamp = rand_datetime()

        subject = external_message["subject"]
        text = external_message["message"]

        if "url" in text:
            text = text.format(url = generate_anchor_tag(f"https://{COMPANY_DOMAIN}/{fake.uri_path()}"))

        message = create_msg(customer, employee, timestamp, subject, text)

        messages.append(message)

    for internal_message in internal_messages:
        sender = select(employees)

        timestamp = rand_datetime()

        subject = internal_message["subject"]
        text = f"Hi {VICTIM["first_name"]},\n\n{internal_message["message"]}\n\nBest,\n{sender["first_name"]}"

        if "url" in text:
            text = text.format(url = generate_anchor_tag(f"https://hq.{COMPANY_DOMAIN}/assets/{rand_int(111111, 999999)}"))

        attachment = None
        if "attach" in text:
            attachment = read_attachment(attachment_index)
            attachment_index += 1
            # num_paragraphs = rand_int(3, 5)
            # attachment = create_docx(fake.paragraphs(nb = num_paragraphs))

        message = create_msg(sender, VICTIM, timestamp, subject, text, attachment)

        messages.append(message)

    return messages


def main(seed : int, messageId : str):
    fake.seed_instance(seed)

    # Generate employees
    for i in range(0, int(len(internal_messages) * INTERNAL_EMAIL_TO_PERSON_RATIO)):
        first_name = fake.first_name()
        last_name = fake.unique.last_name()
        email = f"{first_name[0].lower()}{last_name.lower()}@{COMPANY_DOMAIN}"

        employees.append({
            "first_name" : first_name,
            "last_name" : last_name,
            "email" : email
        })

    # Generate customers
    for i in range(0, int(len(external_messages) * EXTERNAL_EMAIL_TO_PERSON_RATIO)):
        first_name = fake.first_name()
        last_name = fake.unique.last_name()
        email = f"{last_name.lower()}{rand_int(10, 999)}@{fake.free_email_domain()}"

        customers.append({
            "first_name" : first_name,
            "last_name" : last_name,
            "email" : email
        })

    mbox = mailbox.mbox(
        "/data/mailbox.mbox",
        factory=None,
        create=True
    )

    mbox.lock()

    messages = [ ]

    try:
        # Add noise
        messages.extend(generate_noise())
        # Add suspicious email
        sender = select(employees[1:]) # VICTIM is the first entry in `employees`, so we need to omit the first entry

        sus_email = create_msg(
            sender,
            recipient = VICTIM,
            timestamp = rand_datetime(),
            subject = "Updated Timesheets",
            body = f"Hi {VICTIM["first_name"]},\n\nPlease see attached for the updated timesheets.\n\nBest,\n{sender["first_name"]}",
            attachment = payload_attachment_bytes,
            pass_auth = False,
            message_id = f"{messageId}@{COMPANY_DOMAIN}"
        )
        messages.append(sus_email)

        messages = shuffle(messages)
 
        for message in messages:
            mbox.add(message)

        mbox.flush()
    finally:
        mbox.unlock()
        mbox.close()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate passphrases with various options.")
    parser.add_argument("-s", "--seed", type=int, default=DEFAULT_SEED, help=f"Seed value for random number generator (default: {DEFAULT_SEED})")
    parser.add_argument("-m", "--messageId", type=str, help="The Messsage-ID of the malicious email")

    args = parser.parse_args()

    main(**vars(args))