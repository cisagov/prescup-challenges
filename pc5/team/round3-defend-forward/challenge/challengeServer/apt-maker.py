#!/usr/bin/python3

# Copyright 2024 Carnegie Mellon University.
# Released under a MIT (SEI)-style license, please see LICENSE.md in the project 
# root or contact permission@sei.cmu.edu for full terms.


import docx
from docx.shared import Pt
import subprocess
import random

subprocess.run(f"rm -f /home/user/challengeServer/hosted_files/apts.tar.gz", shell=True)
subprocess.run(f"rm -f /home/user/challengeServer/dev/*.pdf", shell=True)
subprocess.run(f"rm -f /home/user/challengeServer/dev/*.docx", shell=True)
subprocess.run(f"rm -f /home/user/challengeServer/dev/*.zip", shell=True)

# Assign Country Name
country_1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_1'", shell=True).decode().strip('\n')
country_2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_2'", shell=True).decode().strip('\n')
country_3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_3'", shell=True).decode().strip('\n')
country_4=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_4'", shell=True).decode().strip('\n')
country_5=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_5'", shell=True).decode().strip('\n')
country_6=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_6'", shell=True).decode().strip('\n')
country_7=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_7'", shell=True).decode().strip('\n')
country_8=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_8'", shell=True).decode().strip('\n')
country_9=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_9'", shell=True).decode().strip('\n')
country_10=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_10'", shell=True).decode().strip('\n')
country_11=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_11'", shell=True).decode().strip('\n')
country_12=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_12'", shell=True).decode().strip('\n')
country_13=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_13'", shell=True).decode().strip('\n')
country_14=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.c_14'", shell=True).decode().strip('\n')

# Assign Population
c1_population=str(random.randint(1000000, 1000000000))
c2_population=str(random.randint(1000000, 1000000000))
c3_population=str(random.randint(1000000, 1000000000))
c4_population=str(random.randint(1000000, 1000000000))
c5_population=str(random.randint(1000000, 1000000000))
c6_population=str(random.randint(1000000, 1000000000))
c7_population=str(random.randint(1000000, 1000000000))
c8_population=str(random.randint(1000000, 1000000000))
c9_population=str(random.randint(1000000, 1000000000))
c10_population=str(random.randint(1000000, 1000000000))
c11_population=str(random.randint(1000000, 1000000000))
c12_population=str(random.randint(1000000, 1000000000))
c13_population=str(random.randint(1000000, 1000000000))
c14_population=str(random.randint(1000000, 1000000000))


# Assign Two Goals
goals = ['Global Civil Unrest', 'Expand Physical Territory', 'Grow Nuclear Capabilities', 'Establish Foothold in Cyberspace', \
	'Access R&D Plans', 'Expand Social Influence', 'Acquire Military Technology', 'Promote Country Beliefs on Others', \
	'Gain Financial Superiority', 'Sabatoge Trust in Rival Governments', 'Control Enemy Critical Infrastructure',]

random.shuffle(goals)
c1_apt1_goal1=goals[0]
c1_apt1_goal2=goals[1]
random.shuffle(goals)
c1_apt2_goal1=goals[0]
c1_apt2_goal2=goals[1]
random.shuffle(goals)
c1_apt3_goal1=goals[0]
c1_apt3_goal2=goals[1]
random.shuffle(goals)
c1_apt4_goal1=goals[0]
c1_apt4_goal2=goals[1]
random.shuffle(goals)
c1_apt5_goal1=goals[0]
c1_apt5_goal2=goals[1]
random.shuffle(goals)
c1_apt6_goal1=goals[0]
c1_apt6_goal2=goals[1]
random.shuffle(goals)
c1_apt7_goal1=goals[0]
c1_apt7_goal2=goals[1]

random.shuffle(goals)
c2_apt1_goal1=goals[0]
c2_apt1_goal2=goals[1]
random.shuffle(goals)
c2_apt2_goal1=goals[0]
c2_apt2_goal2=goals[1]
random.shuffle(goals)
c2_apt3_goal1=goals[0]
c2_apt3_goal2=goals[1]
random.shuffle(goals)
c2_apt4_goal1=goals[0]
c2_apt4_goal2=goals[1]
random.shuffle(goals)
c2_apt5_goal1=goals[0]
c2_apt5_goal2=goals[1]
random.shuffle(goals)
c2_apt6_goal1=goals[0]
c2_apt6_goal2=goals[1]
random.shuffle(goals)
c2_apt7_goal1=goals[0]
c2_apt7_goal2=goals[1]

random.shuffle(goals)
c3_apt1_goal1=goals[0]
c3_apt1_goal2=goals[1]
random.shuffle(goals)
c3_apt2_goal1=goals[0]
c3_apt2_goal2=goals[1]
random.shuffle(goals)
c3_apt3_goal1=goals[0]
c3_apt3_goal2=goals[1]
random.shuffle(goals)
c3_apt4_goal1=goals[0]
c3_apt4_goal2=goals[1]
random.shuffle(goals)
c3_apt5_goal1=goals[0]
c3_apt5_goal2=goals[1]
random.shuffle(goals)
c3_apt6_goal1=goals[0]
c3_apt6_goal2=goals[1]
random.shuffle(goals)
c3_apt7_goal1=goals[0]
c3_apt7_goal2=goals[1]

random.shuffle(goals)
c4_apt1_goal1=goals[0]
c4_apt1_goal2=goals[1]
random.shuffle(goals)
c4_apt2_goal1=goals[0]
c4_apt2_goal2=goals[1]
random.shuffle(goals)
c4_apt3_goal1=goals[0]
c4_apt3_goal2=goals[1]
random.shuffle(goals)
c4_apt4_goal1=goals[0]
c4_apt4_goal2=goals[1]
random.shuffle(goals)
c4_apt5_goal1=goals[0]
c4_apt5_goal2=goals[1]
random.shuffle(goals)
c4_apt6_goal1=goals[0]
c4_apt6_goal2=goals[1]
random.shuffle(goals)
c4_apt7_goal1=goals[0]
c4_apt7_goal2=goals[1]

random.shuffle(goals)
c5_apt1_goal1=goals[0]
c5_apt1_goal2=goals[1]
random.shuffle(goals)
c5_apt2_goal1=goals[0]
c5_apt2_goal2=goals[1]
random.shuffle(goals)
c5_apt3_goal1=goals[0]
c5_apt3_goal2=goals[1]
random.shuffle(goals)
c5_apt4_goal1=goals[0]
c5_apt4_goal2=goals[1]
random.shuffle(goals)
c5_apt5_goal1=goals[0]
c5_apt5_goal2=goals[1]
random.shuffle(goals)
c5_apt6_goal1=goals[0]
c5_apt6_goal2=goals[1]
random.shuffle(goals)
c5_apt7_goal1=goals[0]
c5_apt7_goal2=goals[1]

random.shuffle(goals)
c6_apt1_goal1=goals[0]
c6_apt1_goal2=goals[1]
random.shuffle(goals)
c6_apt2_goal1=goals[0]
c6_apt2_goal2=goals[1]
random.shuffle(goals)
c6_apt3_goal1=goals[0]
c6_apt3_goal2=goals[1]
random.shuffle(goals)
c6_apt4_goal1=goals[0]
c6_apt4_goal2=goals[1]
random.shuffle(goals)
c6_apt5_goal1=goals[0]
c6_apt5_goal2=goals[1]
random.shuffle(goals)
c6_apt6_goal1=goals[0]
c6_apt6_goal2=goals[1]
random.shuffle(goals)
c6_apt7_goal1=goals[0]
c6_apt7_goal2=goals[1]

random.shuffle(goals)
c7_apt1_goal1=goals[0]
c7_apt1_goal2=goals[1]
random.shuffle(goals)
c7_apt2_goal1=goals[0]
c7_apt2_goal2=goals[1]
random.shuffle(goals)
c7_apt3_goal1=goals[0]
c7_apt3_goal2=goals[1]
random.shuffle(goals)
c7_apt4_goal1=goals[0]
c7_apt4_goal2=goals[1]
random.shuffle(goals)
c7_apt5_goal1=goals[0]
c7_apt5_goal2=goals[1]
random.shuffle(goals)
c7_apt6_goal1=goals[0]
c7_apt6_goal2=goals[1]
random.shuffle(goals)
c7_apt7_goal1=goals[0]
c7_apt7_goal2=goals[1]

random.shuffle(goals)
c8_apt1_goal1=goals[0]
c8_apt1_goal2=goals[1]
random.shuffle(goals)
c8_apt2_goal1=goals[0]
c8_apt2_goal2=goals[1]
random.shuffle(goals)
c8_apt3_goal1=goals[0]
c8_apt3_goal2=goals[1]
random.shuffle(goals)
c8_apt4_goal1=goals[0]
c8_apt4_goal2=goals[1]
random.shuffle(goals)
c8_apt5_goal1=goals[0]
c8_apt5_goal2=goals[1]
random.shuffle(goals)
c8_apt6_goal1=goals[0]
c8_apt6_goal2=goals[1]
random.shuffle(goals)
c8_apt7_goal1=goals[0]
c8_apt7_goal2=goals[1]

random.shuffle(goals)
c9_apt1_goal1=goals[0]
c9_apt1_goal2=goals[1]
random.shuffle(goals)
c9_apt2_goal1=goals[0]
c9_apt2_goal2=goals[1]
random.shuffle(goals)
c9_apt3_goal1=goals[0]
c9_apt3_goal2=goals[1]
random.shuffle(goals)
c9_apt4_goal1=goals[0]
c9_apt4_goal2=goals[1]
random.shuffle(goals)
c9_apt5_goal1=goals[0]
c9_apt5_goal2=goals[1]
random.shuffle(goals)
c9_apt6_goal1=goals[0]
c9_apt6_goal2=goals[1]
random.shuffle(goals)
c9_apt7_goal1=goals[0]
c9_apt7_goal2=goals[1]

random.shuffle(goals)
c10_apt1_goal1=goals[0]
c10_apt1_goal2=goals[1]
random.shuffle(goals)
c10_apt2_goal1=goals[0]
c10_apt2_goal2=goals[1]
random.shuffle(goals)
c10_apt3_goal1=goals[0]
c10_apt3_goal2=goals[1]
random.shuffle(goals)
c10_apt4_goal1=goals[0]
c10_apt4_goal2=goals[1]
random.shuffle(goals)
c10_apt5_goal1=goals[0]
c10_apt5_goal2=goals[1]
random.shuffle(goals)
c10_apt6_goal1=goals[0]
c10_apt6_goal2=goals[1]
random.shuffle(goals)
c10_apt7_goal1=goals[0]
c10_apt7_goal2=goals[1]

random.shuffle(goals)
c11_apt1_goal1=goals[0]
c11_apt1_goal2=goals[1]
random.shuffle(goals)
c11_apt2_goal1=goals[0]
c11_apt2_goal2=goals[1]
random.shuffle(goals)
c11_apt3_goal1=goals[0]
c11_apt3_goal2=goals[1]
random.shuffle(goals)
c11_apt4_goal1=goals[0]
c11_apt4_goal2=goals[1]
random.shuffle(goals)
c11_apt5_goal1=goals[0]
c11_apt5_goal2=goals[1]
random.shuffle(goals)
c11_apt6_goal1=goals[0]
c11_apt6_goal2=goals[1]
random.shuffle(goals)
c11_apt7_goal1=goals[0]
c11_apt7_goal2=goals[1]

random.shuffle(goals)
c12_apt1_goal1=goals[0]
c12_apt1_goal2=goals[1]
random.shuffle(goals)
c12_apt2_goal1=goals[0]
c12_apt2_goal2=goals[1]
random.shuffle(goals)
c12_apt3_goal1=goals[0]
c12_apt3_goal2=goals[1]
random.shuffle(goals)
c12_apt4_goal1=goals[0]
c12_apt4_goal2=goals[1]
random.shuffle(goals)
c12_apt5_goal1=goals[0]
c12_apt5_goal2=goals[1]
random.shuffle(goals)
c12_apt6_goal1=goals[0]
c12_apt6_goal2=goals[1]
random.shuffle(goals)
c12_apt7_goal1=goals[0]
c12_apt7_goal2=goals[1]

random.shuffle(goals)
c13_apt1_goal1=goals[0]
c13_apt1_goal2=goals[1]
random.shuffle(goals)
c13_apt2_goal1=goals[0]
c13_apt2_goal2=goals[1]
random.shuffle(goals)
c13_apt3_goal1=goals[0]
c13_apt3_goal2=goals[1]
random.shuffle(goals)
c13_apt4_goal1=goals[0]
c13_apt4_goal2=goals[1]
random.shuffle(goals)
c13_apt5_goal1=goals[0]
c13_apt5_goal2=goals[1]
random.shuffle(goals)
c13_apt6_goal1=goals[0]
c13_apt6_goal2=goals[1]
random.shuffle(goals)
c13_apt7_goal1=goals[0]
c13_apt7_goal2=goals[1]

random.shuffle(goals)
c14_apt1_goal1=goals[0]
c14_apt1_goal2=goals[1]
random.shuffle(goals)
c14_apt2_goal1=goals[0]
c14_apt2_goal2=goals[1]
random.shuffle(goals)
c14_apt3_goal1=goals[0]
c14_apt3_goal2=goals[1]
random.shuffle(goals)
c14_apt4_goal1=goals[0]
c14_apt4_goal2=goals[1]
random.shuffle(goals)
c14_apt5_goal1=goals[0]
c14_apt5_goal2=goals[1]
random.shuffle(goals)
c14_apt6_goal1=goals[0]
c14_apt6_goal2=goals[1]
random.shuffle(goals)
c14_apt7_goal1=goals[0]
c14_apt7_goal2=goals[1]

# Assign 7 APT Group Names
apts_file='apts.txt'
with open(apts_file) as file:
	lines = [line.rstrip() for line in file]
random.shuffle(lines)

c1_apt1=lines[0]
c1_apt2=lines[1]
c1_apt3=lines[2]
c1_apt4=lines[3]
c1_apt5=lines[4]
c1_apt6=lines[5]
c1_apt7=lines[6]

c2_apt1=lines[7]
c2_apt2=lines[8]
c2_apt3=lines[9]
c2_apt4=lines[10]
c2_apt5=lines[11]
c2_apt6=lines[12]
c2_apt7=lines[13]

c3_apt1=lines[14]
c3_apt2=lines[15]
c3_apt3=lines[16]
c3_apt4=lines[17]
c3_apt5=lines[18]
c3_apt6=lines[19]
c3_apt7=lines[20]

c4_apt1=lines[21]
c4_apt2=lines[22]
c4_apt3=lines[23]
c4_apt4=lines[24]
c4_apt5=lines[25]
c4_apt6=lines[26]
c4_apt7=lines[27]

c5_apt1=lines[28]
c5_apt2=lines[29]
c5_apt3=lines[30]
c5_apt4=lines[31]
c5_apt5=lines[32]
c5_apt6=lines[33]
c5_apt7=lines[34]

c6_apt1=lines[35]
c6_apt2=lines[36]
c6_apt3=lines[37]
c6_apt4=lines[38]
c6_apt5=lines[39]
c6_apt6=lines[40]
c6_apt7=lines[41]

c7_apt1=lines[42]
c7_apt2=lines[43]
c7_apt3=lines[44]
c7_apt4=lines[45]
c7_apt5=lines[46]
c7_apt6=lines[47]
c7_apt7=lines[48]

c8_apt1=lines[49]
c8_apt2=lines[50]
c8_apt3=lines[51]
c8_apt4=lines[52]
c8_apt5=lines[53]
c8_apt6=lines[54]
c8_apt7=lines[55]

c9_apt1=lines[56]
c9_apt2=lines[57]
c9_apt3=lines[58]
c9_apt4=lines[59]
c9_apt5=lines[60]
c9_apt6=lines[61]
c9_apt7=lines[62]

c10_apt1=lines[63]
c10_apt2=lines[64]
c10_apt3=lines[65]
c10_apt4=lines[66]
c10_apt5=lines[67]
c10_apt6=lines[68]
c10_apt7=lines[69]

c11_apt1=lines[70]
c11_apt2=lines[71]
c11_apt3=lines[72]
c11_apt4=lines[73]
c11_apt5=lines[74]
c11_apt6=lines[75]
c11_apt7=lines[76]

c12_apt1=lines[77]
c12_apt2=lines[78]
c12_apt3=lines[79]
c12_apt4=lines[80]
c12_apt5=lines[81]
c12_apt6=lines[82]
c12_apt7=lines[83]

c13_apt1=lines[84]
c13_apt2=lines[85]
c13_apt3=lines[86]
c13_apt4=lines[87]
c13_apt5=lines[88]
c13_apt6=lines[89]
c13_apt7=lines[90]

c14_apt1=lines[91]
c14_apt2=lines[92]
c14_apt3=lines[93]
c14_apt4=lines[94]
c14_apt5=lines[95]
c14_apt6=lines[96]
c14_apt7=lines[97]

# Assign Date of First Seen
days=list(range(1, 29))
months='Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec'
years=list(range(2020, 2024))

c1_apt1_day=str(random.choice(days))
c1_apt1_month=random.choice(months)
c1_apt1_year=str(random.choice(years))
c1_apt1_date=c1_apt1_day+c1_apt1_month+c1_apt1_year
print(c1_apt1_date)

c1_apt2_day=str(random.choice(days))
c1_apt2_month=random.choice(months)
c1_apt2_year=str(random.choice(years))
c1_apt2_date=c1_apt2_day+c1_apt2_month+c1_apt2_year
print(c1_apt2_date)

c1_apt3_day=str(random.choice(days))
c1_apt3_month=random.choice(months)
c1_apt3_year=str(random.choice(years))
c1_apt3_date=c1_apt3_day+c1_apt3_month+c1_apt3_year
print(c1_apt3_date)

c1_apt4_day=str(random.choice(days))
c1_apt4_month=random.choice(months)
c1_apt4_year=str(random.choice(years))
c1_apt4_date=c1_apt4_day+c1_apt4_month+c1_apt4_year
print(c1_apt4_date)

c1_apt5_day=str(random.choice(days))
c1_apt5_month=random.choice(months)
c1_apt5_year=str(random.choice(years))
c1_apt5_date=c1_apt5_day+c1_apt5_month+c1_apt5_year
print(c1_apt5_date)

c1_apt6_day=str(random.choice(days))
c1_apt6_month=random.choice(months)
c1_apt6_year=str(random.choice(years))
c1_apt6_date=c1_apt6_day+c1_apt6_month+c1_apt6_year
print(c1_apt6_date)

c1_apt7_day=str(random.choice(days))
c1_apt7_month=random.choice(months)
c1_apt7_year=str(random.choice(years))
c1_apt7_date=c1_apt7_day+c1_apt7_month+c1_apt7_year
print(c1_apt7_date)

c2_apt1_day=str(random.choice(days))
c2_apt1_month=random.choice(months)
c2_apt1_year=str(random.choice(years))
c2_apt1_date=c2_apt1_day+c2_apt1_month+c2_apt1_year
print(c2_apt1_date)

c2_apt2_day=str(random.choice(days))
c2_apt2_month=random.choice(months)
c2_apt2_year=str(random.choice(years))
c2_apt2_date=c2_apt2_day+c2_apt2_month+c2_apt2_year
print(c2_apt2_date)

c2_apt3_day=str(random.choice(days))
c2_apt3_month=random.choice(months)
c2_apt3_year=str(random.choice(years))
c2_apt3_date=c2_apt3_day+c2_apt3_month+c2_apt3_year
print(c2_apt3_date)

c2_apt4_day=str(random.choice(days))
c2_apt4_month=random.choice(months)
c2_apt4_year=str(random.choice(years))
c2_apt4_date=c2_apt4_day+c2_apt4_month+c2_apt4_year
print(c2_apt4_date)

c2_apt5_day=str(random.choice(days))
c2_apt5_month=random.choice(months)
c2_apt5_year=str(random.choice(years))
c2_apt5_date=c2_apt5_day+c2_apt5_month+c2_apt5_year
print(c2_apt5_date)

c2_apt6_day=str(random.choice(days))
c2_apt6_month=random.choice(months)
c2_apt6_year=str(random.choice(years))
c2_apt6_date=c2_apt6_day+c2_apt6_month+c2_apt6_year
print(c2_apt6_date)

c2_apt7_day=str(random.choice(days))
c2_apt7_month=random.choice(months)
c2_apt7_year=str(random.choice(years))
c2_apt7_date=c2_apt7_day+c2_apt7_month+c2_apt7_year
print(c2_apt7_date)
c3_apt1_day=str(random.choice(days))
c3_apt1_month=random.choice(months)
c3_apt1_year=str(random.choice(years))
c3_apt1_date=c3_apt1_day+c3_apt1_month+c3_apt1_year
print(c3_apt1_date)

c3_apt2_day=str(random.choice(days))
c3_apt2_month=random.choice(months)
c3_apt2_year=str(random.choice(years))
c3_apt2_date=c3_apt2_day+c3_apt2_month+c3_apt2_year
print(c3_apt2_date)

c3_apt3_day=str(random.choice(days))
c3_apt3_month=random.choice(months)
c3_apt3_year=str(random.choice(years))
c3_apt3_date=c3_apt3_day+c3_apt3_month+c3_apt3_year
print(c3_apt3_date)

c3_apt4_day=str(random.choice(days))
c3_apt4_month=random.choice(months)
c3_apt4_year=str(random.choice(years))
c3_apt4_date=c3_apt4_day+c3_apt4_month+c3_apt4_year
print(c3_apt4_date)

c3_apt5_day=str(random.choice(days))
c3_apt5_month=random.choice(months)
c3_apt5_year=str(random.choice(years))
c3_apt5_date=c3_apt5_day+c3_apt5_month+c3_apt5_year
print(c3_apt5_date)

c3_apt6_day=str(random.choice(days))
c3_apt6_month=random.choice(months)
c3_apt6_year=str(random.choice(years))
c3_apt6_date=c3_apt6_day+c3_apt6_month+c3_apt6_year
print(c3_apt6_date)

c3_apt7_day=str(random.choice(days))
c3_apt7_month=random.choice(months)
c3_apt7_year=str(random.choice(years))
c3_apt7_date=c3_apt7_day+c3_apt7_month+c3_apt7_year
print(c3_apt7_date)
c4_apt1_day=str(random.choice(days))
c4_apt1_month=random.choice(months)
c4_apt1_year=str(random.choice(years))
c4_apt1_date=c4_apt1_day+c4_apt1_month+c4_apt1_year
print(c4_apt1_date)

c4_apt2_day=str(random.choice(days))
c4_apt2_month=random.choice(months)
c4_apt2_year=str(random.choice(years))
c4_apt2_date=c4_apt2_day+c4_apt2_month+c4_apt2_year
print(c4_apt2_date)

c4_apt3_day=str(random.choice(days))
c4_apt3_month=random.choice(months)
c4_apt3_year=str(random.choice(years))
c4_apt3_date=c4_apt3_day+c4_apt3_month+c4_apt3_year
print(c4_apt3_date)

c4_apt4_day=str(random.choice(days))
c4_apt4_month=random.choice(months)
c4_apt4_year=str(random.choice(years))
c4_apt4_date=c4_apt4_day+c4_apt4_month+c4_apt4_year
print(c4_apt4_date)

c4_apt5_day=str(random.choice(days))
c4_apt5_month=random.choice(months)
c4_apt5_year=str(random.choice(years))
c4_apt5_date=c4_apt5_day+c4_apt5_month+c4_apt5_year
print(c4_apt5_date)

c4_apt6_day=str(random.choice(days))
c4_apt6_month=random.choice(months)
c4_apt6_year=str(random.choice(years))
c4_apt6_date=c4_apt6_day+c4_apt6_month+c4_apt6_year
print(c4_apt6_date)

c4_apt7_day=str(random.choice(days))
c4_apt7_month=random.choice(months)
c4_apt7_year=str(random.choice(years))
c4_apt7_date=c4_apt7_day+c4_apt7_month+c4_apt7_year
print(c4_apt7_date)
c5_apt1_day=str(random.choice(days))
c5_apt1_month=random.choice(months)
c5_apt1_year=str(random.choice(years))
c5_apt1_date=c5_apt1_day+c5_apt1_month+c5_apt1_year
print(c5_apt1_date)

c5_apt2_day=str(random.choice(days))
c5_apt2_month=random.choice(months)
c5_apt2_year=str(random.choice(years))
c5_apt2_date=c5_apt2_day+c5_apt2_month+c5_apt2_year
print(c5_apt2_date)

c5_apt3_day=str(random.choice(days))
c5_apt3_month=random.choice(months)
c5_apt3_year=str(random.choice(years))
c5_apt3_date=c5_apt3_day+c5_apt3_month+c5_apt3_year
print(c5_apt3_date)

c5_apt4_day=str(random.choice(days))
c5_apt4_month=random.choice(months)
c5_apt4_year=str(random.choice(years))
c5_apt4_date=c5_apt4_day+c5_apt4_month+c5_apt4_year
print(c5_apt4_date)

c5_apt5_day=str(random.choice(days))
c5_apt5_month=random.choice(months)
c5_apt5_year=str(random.choice(years))
c5_apt5_date=c5_apt5_day+c5_apt5_month+c5_apt5_year
print(c5_apt5_date)

c5_apt6_day=str(random.choice(days))
c5_apt6_month=random.choice(months)
c5_apt6_year=str(random.choice(years))
c5_apt6_date=c5_apt6_day+c5_apt6_month+c5_apt6_year
print(c5_apt6_date)

c5_apt7_day=str(random.choice(days))
c5_apt7_month=random.choice(months)
c5_apt7_year=str(random.choice(years))
c5_apt7_date=c5_apt7_day+c5_apt7_month+c5_apt7_year
print(c5_apt7_date)
c6_apt1_day=str(random.choice(days))
c6_apt1_month=random.choice(months)
c6_apt1_year=str(random.choice(years))
c6_apt1_date=c6_apt1_day+c6_apt1_month+c6_apt1_year
print(c6_apt1_date)

c6_apt2_day=str(random.choice(days))
c6_apt2_month=random.choice(months)
c6_apt2_year=str(random.choice(years))
c6_apt2_date=c6_apt2_day+c6_apt2_month+c6_apt2_year
print(c6_apt2_date)

c6_apt3_day=str(random.choice(days))
c6_apt3_month=random.choice(months)
c6_apt3_year=str(random.choice(years))
c6_apt3_date=c6_apt3_day+c6_apt3_month+c6_apt3_year
print(c6_apt3_date)

c6_apt4_day=str(random.choice(days))
c6_apt4_month=random.choice(months)
c6_apt4_year=str(random.choice(years))
c6_apt4_date=c6_apt4_day+c6_apt4_month+c6_apt4_year
print(c6_apt4_date)

c6_apt5_day=str(random.choice(days))
c6_apt5_month=random.choice(months)
c6_apt5_year=str(random.choice(years))
c6_apt5_date=c6_apt5_day+c6_apt5_month+c6_apt5_year
print(c6_apt5_date)

c6_apt6_day=str(random.choice(days))
c6_apt6_month=random.choice(months)
c6_apt6_year=str(random.choice(years))
c6_apt6_date=c6_apt6_day+c6_apt6_month+c6_apt6_year
print(c6_apt6_date)

c6_apt7_day=str(random.choice(days))
c6_apt7_month=random.choice(months)
c6_apt7_year=str(random.choice(years))
c6_apt7_date=c6_apt7_day+c6_apt7_month+c6_apt7_year
print(c6_apt7_date)
c7_apt1_day=str(random.choice(days))
c7_apt1_month=random.choice(months)
c7_apt1_year=str(random.choice(years))
c7_apt1_date=c7_apt1_day+c7_apt1_month+c7_apt1_year
print(c7_apt1_date)

c7_apt2_day=str(random.choice(days))
c7_apt2_month=random.choice(months)
c7_apt2_year=str(random.choice(years))
c7_apt2_date=c7_apt2_day+c7_apt2_month+c7_apt2_year
print(c7_apt2_date)

c7_apt3_day=str(random.choice(days))
c7_apt3_month=random.choice(months)
c7_apt3_year=str(random.choice(years))
c7_apt3_date=c7_apt3_day+c7_apt3_month+c7_apt3_year
print(c7_apt3_date)

c7_apt4_day=str(random.choice(days))
c7_apt4_month=random.choice(months)
c7_apt4_year=str(random.choice(years))
c7_apt4_date=c7_apt4_day+c7_apt4_month+c7_apt4_year
print(c7_apt4_date)

c7_apt5_day=str(random.choice(days))
c7_apt5_month=random.choice(months)
c7_apt5_year=str(random.choice(years))
c7_apt5_date=c7_apt5_day+c7_apt5_month+c7_apt5_year
print(c7_apt5_date)

c7_apt6_day=str(random.choice(days))
c7_apt6_month=random.choice(months)
c7_apt6_year=str(random.choice(years))
c7_apt6_date=c7_apt6_day+c7_apt6_month+c7_apt6_year
print(c7_apt6_date)

c7_apt7_day=str(random.choice(days))
c7_apt7_month=random.choice(months)
c7_apt7_year=str(random.choice(years))
c7_apt7_date=c7_apt7_day+c7_apt7_month+c7_apt7_year
print(c7_apt7_date)

c8_apt1_day=str(random.choice(days))
c8_apt1_month=random.choice(months)
c8_apt1_year=str(random.choice(years))
c8_apt1_date=c8_apt1_day+c8_apt1_month+c8_apt1_year
print(c8_apt1_date)

c8_apt2_day=str(random.choice(days))
c8_apt2_month=random.choice(months)
c8_apt2_year=str(random.choice(years))
c8_apt2_date=c8_apt2_day+c8_apt2_month+c8_apt2_year
print(c8_apt2_date)

c8_apt3_day=str(random.choice(days))
c8_apt3_month=random.choice(months)
c8_apt3_year=str(random.choice(years))
c8_apt3_date=c8_apt3_day+c8_apt3_month+c8_apt3_year
print(c8_apt3_date)

c8_apt4_day=str(random.choice(days))
c8_apt4_month=random.choice(months)
c8_apt4_year=str(random.choice(years))
c8_apt4_date=c8_apt4_day+c8_apt4_month+c8_apt4_year
print(c8_apt4_date)

c8_apt5_day=str(random.choice(days))
c8_apt5_month=random.choice(months)
c8_apt5_year=str(random.choice(years))
c8_apt5_date=c8_apt5_day+c8_apt5_month+c8_apt5_year
print(c8_apt5_date)

c8_apt6_day=str(random.choice(days))
c8_apt6_month=random.choice(months)
c8_apt6_year=str(random.choice(years))
c8_apt6_date=c8_apt6_day+c8_apt6_month+c8_apt6_year
print(c8_apt6_date)

c8_apt7_day=str(random.choice(days))
c8_apt7_month=random.choice(months)
c8_apt7_year=str(random.choice(years))
c8_apt7_date=c8_apt7_day+c8_apt7_month+c8_apt7_year
print(c8_apt7_date)

c9_apt1_day=str(random.choice(days))
c9_apt1_month=random.choice(months)
c9_apt1_year=str(random.choice(years))
c9_apt1_date=c9_apt1_day+c9_apt1_month+c9_apt1_year
print(c9_apt1_date)

c9_apt2_day=str(random.choice(days))
c9_apt2_month=random.choice(months)
c9_apt2_year=str(random.choice(years))
c9_apt2_date=c9_apt2_day+c9_apt2_month+c9_apt2_year
print(c9_apt2_date)

c9_apt3_day=str(random.choice(days))
c9_apt3_month=random.choice(months)
c9_apt3_year=str(random.choice(years))
c9_apt3_date=c9_apt3_day+c9_apt3_month+c9_apt3_year
print(c9_apt3_date)

c9_apt4_day=str(random.choice(days))
c9_apt4_month=random.choice(months)
c9_apt4_year=str(random.choice(years))
c9_apt4_date=c9_apt4_day+c9_apt4_month+c9_apt4_year
print(c9_apt4_date)

c9_apt5_day=str(random.choice(days))
c9_apt5_month=random.choice(months)
c9_apt5_year=str(random.choice(years))
c9_apt5_date=c9_apt5_day+c9_apt5_month+c9_apt5_year
print(c9_apt5_date)

c9_apt6_day=str(random.choice(days))
c9_apt6_month=random.choice(months)
c9_apt6_year=str(random.choice(years))
c9_apt6_date=c9_apt6_day+c9_apt6_month+c9_apt6_year
print(c9_apt6_date)

c9_apt7_day=str(random.choice(days))
c9_apt7_month=random.choice(months)
c9_apt7_year=str(random.choice(years))
c9_apt7_date=c9_apt7_day+c9_apt7_month+c9_apt7_year
print(c9_apt7_date)
c10_apt1_day=str(random.choice(days))
c10_apt1_month=random.choice(months)
c10_apt1_year=str(random.choice(years))
c10_apt1_date=c10_apt1_day+c10_apt1_month+c10_apt1_year
print(c10_apt1_date)

c10_apt2_day=str(random.choice(days))
c10_apt2_month=random.choice(months)
c10_apt2_year=str(random.choice(years))
c10_apt2_date=c10_apt2_day+c10_apt2_month+c10_apt2_year
print(c10_apt2_date)

c10_apt3_day=str(random.choice(days))
c10_apt3_month=random.choice(months)
c10_apt3_year=str(random.choice(years))
c10_apt3_date=c10_apt3_day+c10_apt3_month+c10_apt3_year
print(c10_apt3_date)

c10_apt4_day=str(random.choice(days))
c10_apt4_month=random.choice(months)
c10_apt4_year=str(random.choice(years))
c10_apt4_date=c10_apt4_day+c10_apt4_month+c10_apt4_year
print(c10_apt4_date)

c10_apt5_day=str(random.choice(days))
c10_apt5_month=random.choice(months)
c10_apt5_year=str(random.choice(years))
c10_apt5_date=c10_apt5_day+c10_apt5_month+c10_apt5_year
print(c10_apt5_date)

c10_apt6_day=str(random.choice(days))
c10_apt6_month=random.choice(months)
c10_apt6_year=str(random.choice(years))
c10_apt6_date=c10_apt6_day+c10_apt6_month+c10_apt6_year
print(c10_apt6_date)

c10_apt7_day=str(random.choice(days))
c10_apt7_month=random.choice(months)
c10_apt7_year=str(random.choice(years))
c10_apt7_date=c10_apt7_day+c10_apt7_month+c10_apt7_year
print(c10_apt7_date)
c11_apt1_day=str(random.choice(days))
c11_apt1_month=random.choice(months)
c11_apt1_year=str(random.choice(years))
c11_apt1_date=c11_apt1_day+c11_apt1_month+c11_apt1_year
print(c11_apt1_date)

c11_apt2_day=str(random.choice(days))
c11_apt2_month=random.choice(months)
c11_apt2_year=str(random.choice(years))
c11_apt2_date=c11_apt2_day+c11_apt2_month+c11_apt2_year
print(c11_apt2_date)

c11_apt3_day=str(random.choice(days))
c11_apt3_month=random.choice(months)
c11_apt3_year=str(random.choice(years))
c11_apt3_date=c11_apt3_day+c11_apt3_month+c11_apt3_year
print(c11_apt3_date)

c11_apt4_day=str(random.choice(days))
c11_apt4_month=random.choice(months)
c11_apt4_year=str(random.choice(years))
c11_apt4_date=c11_apt4_day+c11_apt4_month+c11_apt4_year
print(c11_apt4_date)

c11_apt5_day=str(random.choice(days))
c11_apt5_month=random.choice(months)
c11_apt5_year=str(random.choice(years))
c11_apt5_date=c11_apt5_day+c11_apt5_month+c11_apt5_year
print(c11_apt5_date)

c11_apt6_day=str(random.choice(days))
c11_apt6_month=random.choice(months)
c11_apt6_year=str(random.choice(years))
c11_apt6_date=c11_apt6_day+c11_apt6_month+c11_apt6_year
print(c11_apt6_date)

c11_apt7_day=str(random.choice(days))
c11_apt7_month=random.choice(months)
c11_apt7_year=str(random.choice(years))
c11_apt7_date=c11_apt7_day+c11_apt7_month+c11_apt7_year
print(c11_apt7_date)
c12_apt1_day=str(random.choice(days))
c12_apt1_month=random.choice(months)
c12_apt1_year=str(random.choice(years))
c12_apt1_date=c12_apt1_day+c12_apt1_month+c12_apt1_year
print(c12_apt1_date)

c12_apt2_day=str(random.choice(days))
c12_apt2_month=random.choice(months)
c12_apt2_year=str(random.choice(years))
c12_apt2_date=c12_apt2_day+c12_apt2_month+c12_apt2_year
print(c12_apt2_date)

c12_apt3_day=str(random.choice(days))
c12_apt3_month=random.choice(months)
c12_apt3_year=str(random.choice(years))
c12_apt3_date=c12_apt3_day+c12_apt3_month+c12_apt3_year
print(c12_apt3_date)

c12_apt4_day=str(random.choice(days))
c12_apt4_month=random.choice(months)
c12_apt4_year=str(random.choice(years))
c12_apt4_date=c12_apt4_day+c12_apt4_month+c12_apt4_year
print(c12_apt4_date)

c12_apt5_day=str(random.choice(days))
c12_apt5_month=random.choice(months)
c12_apt5_year=str(random.choice(years))
c12_apt5_date=c12_apt5_day+c12_apt5_month+c12_apt5_year
print(c12_apt5_date)

c12_apt6_day=str(random.choice(days))
c12_apt6_month=random.choice(months)
c12_apt6_year=str(random.choice(years))
c12_apt6_date=c12_apt6_day+c12_apt6_month+c12_apt6_year
print(c12_apt6_date)

c12_apt7_day=str(random.choice(days))
c12_apt7_month=random.choice(months)
c12_apt7_year=str(random.choice(years))
c12_apt7_date=c12_apt7_day+c12_apt7_month+c12_apt7_year
print(c12_apt7_date)
c13_apt1_day=str(random.choice(days))
c13_apt1_month=random.choice(months)
c13_apt1_year=str(random.choice(years))
c13_apt1_date=c13_apt1_day+c13_apt1_month+c13_apt1_year
print(c13_apt1_date)

c13_apt2_day=str(random.choice(days))
c13_apt2_month=random.choice(months)
c13_apt2_year=str(random.choice(years))
c13_apt2_date=c13_apt2_day+c13_apt2_month+c13_apt2_year
print(c13_apt2_date)

c13_apt3_day=str(random.choice(days))
c13_apt3_month=random.choice(months)
c13_apt3_year=str(random.choice(years))
c13_apt3_date=c13_apt3_day+c13_apt3_month+c13_apt3_year
print(c13_apt3_date)

c13_apt4_day=str(random.choice(days))
c13_apt4_month=random.choice(months)
c13_apt4_year=str(random.choice(years))
c13_apt4_date=c13_apt4_day+c13_apt4_month+c13_apt4_year
print(c13_apt4_date)

c13_apt5_day=str(random.choice(days))
c13_apt5_month=random.choice(months)
c13_apt5_year=str(random.choice(years))
c13_apt5_date=c13_apt5_day+c13_apt5_month+c13_apt5_year
print(c13_apt5_date)

c13_apt6_day=str(random.choice(days))
c13_apt6_month=random.choice(months)
c13_apt6_year=str(random.choice(years))
c13_apt6_date=c13_apt6_day+c13_apt6_month+c13_apt6_year
print(c13_apt6_date)

c13_apt7_day=str(random.choice(days))
c13_apt7_month=random.choice(months)
c13_apt7_year=str(random.choice(years))
c13_apt7_date=c13_apt7_day+c13_apt7_month+c13_apt7_year
print(c13_apt7_date)
c14_apt1_day=str(random.choice(days))
c14_apt1_month=random.choice(months)
c14_apt1_year=str(random.choice(years))
c14_apt1_date=c14_apt1_day+c14_apt1_month+c14_apt1_year
print(c14_apt1_date)

c14_apt2_day=str(random.choice(days))
c14_apt2_month=random.choice(months)
c14_apt2_year=str(random.choice(years))
c14_apt2_date=c14_apt2_day+c14_apt2_month+c14_apt2_year
print(c14_apt2_date)

c14_apt3_day=str(random.choice(days))
c14_apt3_month=random.choice(months)
c14_apt3_year=str(random.choice(years))
c14_apt3_date=c14_apt3_day+c14_apt3_month+c14_apt3_year
print(c14_apt3_date)

c14_apt4_day=str(random.choice(days))
c14_apt4_month=random.choice(months)
c14_apt4_year=str(random.choice(years))
c14_apt4_date=c14_apt4_day+c14_apt4_month+c14_apt4_year
print(c14_apt4_date)

c14_apt5_day=str(random.choice(days))
c14_apt5_month=random.choice(months)
c14_apt5_year=str(random.choice(years))
c14_apt5_date=c14_apt5_day+c14_apt5_month+c14_apt5_year
print(c14_apt5_date)

c14_apt6_day=str(random.choice(days))
c14_apt6_month=random.choice(months)
c14_apt6_year=str(random.choice(years))
c14_apt6_date=c14_apt6_day+c14_apt6_month+c14_apt6_year
print(c14_apt6_date)

c14_apt7_day=str(random.choice(days))
c14_apt7_month=random.choice(months)
c14_apt7_year=str(random.choice(years))
c14_apt7_date=c14_apt7_day+c14_apt7_month+c14_apt7_year
print(c14_apt7_date)

# Assign 7 Targets to Hurt APT Group. c1_apt1, c2_apt1, and c3_apt1 are the targets.

tgts = ['Hacker Group/Government Website', 'Database', 'Bot/C2 Machine', 'Router', 'SCADA', 'Country Firewall', 'GPS Systems']

c1_apt1_target=tgts[0]
c1_apt2_target=tgts[1]
c1_apt3_target=tgts[2]
c1_apt4_target=tgts[3]
c1_apt5_target=tgts[4]
c1_apt6_target=tgts[5]
c1_apt7_target=tgts[6]

c2_apt1_target=tgts[1]
c2_apt2_target=tgts[0]
c2_apt3_target=tgts[2]
c2_apt4_target=tgts[3]
c2_apt5_target=tgts[4]
c2_apt6_target=tgts[5]
c2_apt7_target=tgts[6]

c3_apt1_target=tgts[2]
c3_apt2_target=tgts[1]
c3_apt3_target=tgts[0]
c3_apt4_target=tgts[3]
c3_apt5_target=tgts[4]
c3_apt6_target=tgts[5]
c3_apt7_target=tgts[6]

random.shuffle(tgts)
c4_apt1_target=tgts[0]
c4_apt2_target=tgts[1]
c4_apt3_target=tgts[2]
c4_apt4_target=tgts[3]
c4_apt5_target=tgts[4]
c4_apt6_target=tgts[5]
c4_apt7_target=tgts[6]

random.shuffle(tgts)
c5_apt1_target=tgts[0]
c5_apt2_target=tgts[1]
c5_apt3_target=tgts[2]
c5_apt4_target=tgts[3]
c5_apt5_target=tgts[4]
c5_apt6_target=tgts[5]
c5_apt7_target=tgts[6]

random.shuffle(tgts)
c6_apt1_target=tgts[0]
c6_apt2_target=tgts[1]
c6_apt3_target=tgts[2]
c6_apt4_target=tgts[3]
c6_apt5_target=tgts[4]
c6_apt6_target=tgts[5]
c6_apt7_target=tgts[6]

random.shuffle(tgts)
c7_apt1_target=tgts[0]
c7_apt2_target=tgts[1]
c7_apt3_target=tgts[2]
c7_apt4_target=tgts[3]
c7_apt5_target=tgts[4]
c7_apt6_target=tgts[5]
c7_apt7_target=tgts[6]

random.shuffle(tgts)
c8_apt1_target=tgts[0]
c8_apt2_target=tgts[1]
c8_apt3_target=tgts[2]
c8_apt4_target=tgts[3]
c8_apt5_target=tgts[4]
c8_apt6_target=tgts[5]
c8_apt7_target=tgts[6]

random.shuffle(tgts)
c9_apt1_target=tgts[0]
c9_apt2_target=tgts[1]
c9_apt3_target=tgts[2]
c9_apt4_target=tgts[3]
c9_apt5_target=tgts[4]
c9_apt6_target=tgts[5]
c9_apt7_target=tgts[6]

random.shuffle(tgts)
c10_apt1_target=tgts[0]
c10_apt2_target=tgts[1]
c10_apt3_target=tgts[2]
c10_apt4_target=tgts[3]
c10_apt5_target=tgts[4]
c10_apt6_target=tgts[5]
c10_apt7_target=tgts[6]

random.shuffle(tgts)
c11_apt1_target=tgts[0]
c11_apt2_target=tgts[1]
c11_apt3_target=tgts[2]
c11_apt4_target=tgts[3]
c11_apt5_target=tgts[4]
c11_apt6_target=tgts[5]
c11_apt7_target=tgts[6]

random.shuffle(tgts)
c12_apt1_target=tgts[0]
c12_apt2_target=tgts[1]
c12_apt3_target=tgts[2]
c12_apt4_target=tgts[3]
c12_apt5_target=tgts[4]
c12_apt6_target=tgts[5]
c12_apt7_target=tgts[6]

random.shuffle(tgts)
c13_apt1_target=tgts[0]
c13_apt2_target=tgts[1]
c13_apt3_target=tgts[2]
c13_apt4_target=tgts[3]
c13_apt5_target=tgts[4]
c13_apt6_target=tgts[5]
c13_apt7_target=tgts[6]

random.shuffle(tgts)
c14_apt1_target=tgts[0]
c14_apt2_target=tgts[1]
c14_apt3_target=tgts[2]
c14_apt4_target=tgts[3]
c14_apt5_target=tgts[4]
c14_apt6_target=tgts[5]
c14_apt7_target=tgts[6]


# Country IP Ranges

c1_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f1", shell=True).decode().strip('\n')
c1_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f1", shell=True).decode().strip('\n')
c1_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f1", shell=True).decode().strip('\n')
c1_ip=str(c1_octet1)+'.'+str(c1_octet2)+'.'+str(c1_octet3)+'.0/24'
print(c1_ip)

c2_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f2", shell=True).decode().strip('\n')
c2_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f2", shell=True).decode().strip('\n')
c2_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f2", shell=True).decode().strip('\n')
c2_ip=str(c2_octet1)+'.'+str(c2_octet2)+'.'+str(c2_octet3)+'.0/24'
print(c2_ip)

c3_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f3", shell=True).decode().strip('\n')
c3_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f3", shell=True).decode().strip('\n')
c3_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f3", shell=True).decode().strip('\n')
c3_ip=str(c3_octet1)+'.'+str(c3_octet2)+'.'+str(c3_octet3)+'.0/24'
print(c3_ip)

c4_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f4", shell=True).decode().strip('\n')
c4_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f4", shell=True).decode().strip('\n')
c4_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f4", shell=True).decode().strip('\n')
c4_ip=str(c4_octet1)+'.'+str(c4_octet2)+'.'+str(c4_octet3)+'.0/24'
print(c4_ip)

c5_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f5", shell=True).decode().strip('\n')
c5_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f5", shell=True).decode().strip('\n')
c5_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f5", shell=True).decode().strip('\n')
c5_ip=str(c5_octet1)+'.'+str(c5_octet2)+'.'+str(c5_octet3)+'.0/24'
print(c5_ip)

c6_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f6", shell=True).decode().strip('\n')
c6_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f6", shell=True).decode().strip('\n')
c6_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f6", shell=True).decode().strip('\n')
c6_ip=str(c6_octet1)+'.'+str(c6_octet2)+'.'+str(c6_octet3)+'.0/24'
print(c6_ip)

c7_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f7", shell=True).decode().strip('\n')
c7_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f7", shell=True).decode().strip('\n')
c7_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f7", shell=True).decode().strip('\n')
c7_ip=str(c7_octet1)+'.'+str(c7_octet2)+'.'+str(c7_octet3)+'.0/24'
print(c7_ip)

c8_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f8", shell=True).decode().strip('\n')
c8_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f8", shell=True).decode().strip('\n')
c8_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f8", shell=True).decode().strip('\n')
c8_ip=str(c8_octet1)+'.'+str(c8_octet2)+'.'+str(c8_octet3)+'.0/24'
print(c8_ip)

c9_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f9", shell=True).decode().strip('\n')
c9_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f9", shell=True).decode().strip('\n')
c9_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f9", shell=True).decode().strip('\n')
c9_ip=str(c9_octet1)+'.'+str(c9_octet2)+'.'+str(c9_octet3)+'.0/24'
print(c9_ip)

c10_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f10", shell=True).decode().strip('\n')
c10_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f10", shell=True).decode().strip('\n')
c10_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f10", shell=True).decode().strip('\n')
c10_ip=str(c10_octet1)+'.'+str(c10_octet2)+'.'+str(c10_octet3)+'.0/24'
print(c10_ip)

c11_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f11", shell=True).decode().strip('\n')
c11_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f11", shell=True).decode().strip('\n')
c11_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f11", shell=True).decode().strip('\n')
c11_ip=str(c11_octet1)+'.'+str(c11_octet2)+'.'+str(c11_octet3)+'.0/24'
print(c11_ip)

c12_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f12", shell=True).decode().strip('\n')
c12_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f12", shell=True).decode().strip('\n')
c12_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f12", shell=True).decode().strip('\n')
c12_ip=str(c12_octet1)+'.'+str(c12_octet2)+'.'+str(c12_octet3)+'.0/24'
print(c12_ip)

c13_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f13", shell=True).decode().strip('\n')
c13_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f13", shell=True).decode().strip('\n')
c13_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f13", shell=True).decode().strip('\n')
c13_ip=str(c13_octet1)+'.'+str(c13_octet2)+'.'+str(c13_octet3)+'.0/24'
print(c13_ip)

c14_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f14", shell=True).decode().strip('\n')
c14_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f14", shell=True).decode().strip('\n')
c14_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f14", shell=True).decode().strip('\n')
c14_ip=str(c14_octet1)+'.'+str(c14_octet2)+'.'+str(c14_octet3)+'.0/24'
print(c14_ip)

def write_plaintext(country, population, apt1, apt2, apt3, apt4, apt5, apt6, apt7, filename):
    file=country+'.docx'
    doc = docx.Document()
    doc.add_heading(country, 0)
    para = doc.add_paragraph().add_run(
            'Population: '+population)
    para = doc.add_paragraph().add_run(
            'Suspected Critical IP Block: THIS WILL BE PROVIDED IF YOU FIND THE THREE ATTACKER IPs and SUBMIT THEM TO challenge.us')
    para = doc.add_paragraph().add_run(
            'Suspected APT Groups: \n'+apt1+'\n'+apt2+'\n'+apt3+'\n'+apt4+'\n'+apt5+'\n'+apt6+'\n'+apt7+'\n')
    para.font.size = Pt(12)
    doc.save(file)
    output = subprocess.check_output(['libreoffice', '--convert-to', 'pdf', file])

write_plaintext(country=country_1, population=c1_population, apt1=c1_apt1, apt2=c1_apt2, apt3=c1_apt3, apt4=c1_apt4, apt5=c1_apt5, apt6=c1_apt6, apt7=c1_apt7, filename=country_1)
write_plaintext(country=country_2, population=c2_population, apt1=c2_apt1, apt2=c2_apt2, apt3=c2_apt3, apt4=c2_apt4, apt5=c2_apt5, apt6=c2_apt6, apt7=c2_apt7, filename=country_2)
write_plaintext(country=country_3, population=c3_population, apt1=c3_apt1, apt2=c3_apt2, apt3=c3_apt3, apt4=c3_apt4, apt5=c3_apt5, apt6=c3_apt6, apt7=c3_apt7, filename=country_3)
write_plaintext(country=country_4, population=c4_population, apt1=c4_apt1, apt2=c4_apt2, apt3=c4_apt3, apt4=c4_apt4, apt5=c4_apt5, apt6=c4_apt6, apt7=c4_apt7, filename=country_4)
write_plaintext(country=country_5, population=c5_population, apt1=c5_apt1, apt2=c5_apt2, apt3=c5_apt3, apt4=c5_apt4, apt5=c5_apt5, apt6=c5_apt6, apt7=c5_apt7, filename=country_5)
write_plaintext(country=country_6, population=c6_population, apt1=c6_apt1, apt2=c6_apt2, apt3=c6_apt3, apt4=c6_apt4, apt5=c6_apt5, apt6=c6_apt6, apt7=c6_apt7, filename=country_6)
write_plaintext(country=country_7, population=c7_population, apt1=c7_apt1, apt2=c7_apt2, apt3=c7_apt3, apt4=c7_apt4, apt5=c7_apt5, apt6=c7_apt6, apt7=c7_apt7, filename=country_7)
write_plaintext(country=country_8, population=c8_population, apt1=c8_apt1, apt2=c8_apt2, apt3=c8_apt3, apt4=c8_apt4, apt5=c8_apt5, apt6=c8_apt6, apt7=c8_apt7, filename=country_8)
write_plaintext(country=country_9, population=c9_population, apt1=c9_apt1, apt2=c9_apt2, apt3=c9_apt3, apt4=c9_apt4, apt5=c9_apt5, apt6=c9_apt6, apt7=c9_apt7, filename=country_9)
write_plaintext(country=country_10, population=c10_population, apt1=c10_apt1, apt2=c10_apt2, apt3=c10_apt3, apt4=c10_apt4, apt5=c10_apt5, apt6=c10_apt6, apt7=c10_apt7, filename=country_10)
write_plaintext(country=country_11, population=c11_population, apt1=c11_apt1, apt2=c11_apt2, apt3=c11_apt3, apt4=c11_apt4, apt5=c11_apt5, apt6=c11_apt6, apt7=c11_apt7, filename=country_11)
write_plaintext(country=country_12, population=c12_population, apt1=c12_apt1, apt2=c12_apt2, apt3=c12_apt3, apt4=c12_apt4, apt5=c12_apt5, apt6=c12_apt6, apt7=c12_apt7, filename=country_12)
write_plaintext(country=country_13, population=c13_population, apt1=c13_apt1, apt2=c13_apt2, apt3=c13_apt3, apt4=c13_apt4, apt5=c13_apt5, apt6=c13_apt6, apt7=c13_apt7, filename=country_13)
write_plaintext(country=country_14, population=c14_population, apt1=c14_apt1, apt2=c14_apt2, apt3=c14_apt3, apt4=c14_apt4, apt5=c14_apt5, apt6=c14_apt6, apt7=c14_apt7, filename=country_14)


def write(country, population, ip, apt1, apt2, apt3, apt4, apt5, apt6, apt7, filename):
    file="ips-"+country+".docx"
    doc = docx.Document()
    doc.add_heading(country, 0)
    para = doc.add_paragraph().add_run(
            'Population: '+population)
    para = doc.add_paragraph().add_run(
            'Suspected Critical IP Block: '+ip)
    para = doc.add_paragraph().add_run(
            'Suspected APT Groups: \n'+apt1+'\n'+apt2+'\n'+apt3+'\n'+apt4+'\n'+apt5+'\n'+apt6+'\n'+apt7+'\n')
    para.font.size = Pt(12)
    doc.save(file)
    output = subprocess.check_output(['libreoffice', '--convert-to', 'pdf', file])

write(country=country_1, population=c1_population, ip=c1_ip, apt1=c1_apt1, apt2=c1_apt2, apt3=c1_apt3, apt4=c1_apt4, apt5=c1_apt5, apt6=c1_apt6, apt7=c1_apt7, filename=country_1)
write(country=country_2, population=c2_population, ip=c2_ip, apt1=c2_apt1, apt2=c2_apt2, apt3=c2_apt3, apt4=c2_apt4, apt5=c2_apt5, apt6=c2_apt6, apt7=c2_apt7, filename=country_2)
write(country=country_3, population=c3_population, ip=c3_ip, apt1=c3_apt1, apt2=c3_apt2, apt3=c3_apt3, apt4=c3_apt4, apt5=c3_apt5, apt6=c3_apt6, apt7=c3_apt7, filename=country_3)
write(country=country_4, population=c4_population, ip=c4_ip, apt1=c4_apt1, apt2=c4_apt2, apt3=c4_apt3, apt4=c4_apt4, apt5=c4_apt5, apt6=c4_apt6, apt7=c4_apt7, filename=country_4)
write(country=country_5, population=c5_population, ip=c5_ip, apt1=c5_apt1, apt2=c5_apt2, apt3=c5_apt3, apt4=c5_apt4, apt5=c5_apt5, apt6=c5_apt6, apt7=c5_apt7, filename=country_5)
write(country=country_6, population=c6_population, ip=c6_ip, apt1=c6_apt1, apt2=c6_apt2, apt3=c6_apt3, apt4=c6_apt4, apt5=c6_apt5, apt6=c6_apt6, apt7=c6_apt7, filename=country_6)
write(country=country_7, population=c7_population, ip=c7_ip, apt1=c7_apt1, apt2=c7_apt2, apt3=c7_apt3, apt4=c7_apt4, apt5=c7_apt5, apt6=c7_apt6, apt7=c7_apt7, filename=country_7)
write(country=country_8, population=c8_population, ip=c8_ip, apt1=c8_apt1, apt2=c8_apt2, apt3=c8_apt3, apt4=c8_apt4, apt5=c8_apt5, apt6=c8_apt6, apt7=c8_apt7, filename=country_8)
write(country=country_9, population=c9_population, ip=c9_ip, apt1=c9_apt1, apt2=c9_apt2, apt3=c9_apt3, apt4=c9_apt4, apt5=c9_apt5, apt6=c9_apt6, apt7=c9_apt7, filename=country_9)
write(country=country_10, population=c10_population, ip=c10_ip, apt1=c10_apt1, apt2=c10_apt2, apt3=c10_apt3, apt4=c10_apt4, apt5=c10_apt5, apt6=c10_apt6, apt7=c10_apt7, filename=country_10)
write(country=country_11, population=c11_population, ip=c11_ip, apt1=c11_apt1, apt2=c11_apt2, apt3=c11_apt3, apt4=c11_apt4, apt5=c11_apt5, apt6=c11_apt6, apt7=c11_apt7, filename=country_11)
write(country=country_12, population=c12_population, ip=c12_ip, apt1=c12_apt1, apt2=c12_apt2, apt3=c12_apt3, apt4=c12_apt4, apt5=c12_apt5, apt6=c12_apt6, apt7=c12_apt7, filename=country_12)
write(country=country_13, population=c13_population, ip=c13_ip, apt1=c13_apt1, apt2=c13_apt2, apt3=c13_apt3, apt4=c13_apt4, apt5=c13_apt5, apt6=c13_apt6, apt7=c13_apt7, filename=country_13)
write(country=country_14, population=c14_population, ip=c14_ip, apt1=c14_apt1, apt2=c14_apt2, apt3=c14_apt3, apt4=c14_apt4, apt5=c14_apt5, apt6=c14_apt6, apt7=c14_apt7, filename=country_14)

#Attacker 1 IP Range
a1_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f12", shell=True).decode().strip('\n')
a1_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f15", shell=True).decode().strip('\n')
a1_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f15", shell=True).decode().strip('\n')
a1_ip=str(a1_octet1)+'.'+str(a1_octet2)+'.'+str(a1_octet3)+'.0/24'
print(a1_ip)

#Attacker 2 IP Range
a2_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f13", shell=True).decode().strip('\n')
a2_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f16", shell=True).decode().strip('\n')
a2_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f16", shell=True).decode().strip('\n')
a2_ip=str(a2_octet1)+'.'+str(a2_octet2)+'.'+str(a2_octet3)+'.0/24'
print(a2_ip)

#Attacker 3 IP Range
a3_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f14", shell=True).decode().strip('\n')
a3_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f17", shell=True).decode().strip('\n')
a3_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f17", shell=True).decode().strip('\n')
a3_ip=str(a3_octet1)+'.'+str(a3_octet2)+'.'+str(a3_octet3)+'.0/24'
print(a3_ip)

#Fluff 1 IP Range
f1_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f2", shell=True).decode().strip('\n')
f1_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f18", shell=True).decode().strip('\n')
f1_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f18", shell=True).decode().strip('\n')
f1_ip=str(f1_octet1)+'.'+str(f1_octet2)+'.'+str(f1_octet3)+'.0/24'
print(f1_ip)

#Fluff 2 IP Range
f2_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f3", shell=True).decode().strip('\n')
f2_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f19", shell=True).decode().strip('\n')
f2_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f19", shell=True).decode().strip('\n')
f2_ip=str(f2_octet1)+'.'+str(f2_octet2)+'.'+str(f2_octet3)+'.0/24'
print(f2_ip)

#Fluff 3 IP Range
f3_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f4", shell=True).decode().strip('\n')
f3_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f20", shell=True).decode().strip('\n')
f3_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f20", shell=True).decode().strip('\n')
f3_ip=str(f3_octet1)+'.'+str(f3_octet2)+'.'+str(f3_octet3)+'.0/24'
print(f3_ip)

#Fluff 4 IP Range
f4_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f5", shell=True).decode().strip('\n')
f4_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f21", shell=True).decode().strip('\n')
f4_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f21", shell=True).decode().strip('\n')
f4_ip=str(f4_octet1)+'.'+str(f4_octet2)+'.'+str(f4_octet3)+'.0/24'
print(f4_ip)

#Fluff 5 IP Range
f5_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f6", shell=True).decode().strip('\n')
f5_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f22", shell=True).decode().strip('\n')
f5_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f22", shell=True).decode().strip('\n')
f5_ip=str(f5_octet1)+'.'+str(f5_octet2)+'.'+str(f5_octet3)+'.0/24'
print(f5_ip)

#Fluff 6 IP Range
f6_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f7", shell=True).decode().strip('\n')
f6_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f23", shell=True).decode().strip('\n')
f6_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f23", shell=True).decode().strip('\n')
f6_ip=str(f6_octet1)+'.'+str(f6_octet2)+'.'+str(f6_octet3)+'.0/24'
print(f6_ip)

#Fluff 7 IP Range
f7_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f8", shell=True).decode().strip('\n')
f7_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f24", shell=True).decode().strip('\n')
f7_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f24", shell=True).decode().strip('\n')
f7_ip=str(f7_octet1)+'.'+str(f7_octet2)+'.'+str(f7_octet3)+'.0/24'
print(f7_ip)

#Fluff 8 IP Range
f8_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f9", shell=True).decode().strip('\n')
f8_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f25", shell=True).decode().strip('\n')
f8_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f25", shell=True).decode().strip('\n')
f8_ip=str(f8_octet1)+'.'+str(f8_octet2)+'.'+str(f8_octet3)+'.0/24'
print(f8_ip)

#Fluff 9 IP Range
f9_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f10", shell=True).decode().strip('\n')
f9_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f26", shell=True).decode().strip('\n')
f9_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f26", shell=True).decode().strip('\n')
f9_ip=str(f9_octet1)+'.'+str(f9_octet2)+'.'+str(f9_octet3)+'.0/24'
print(f9_ip)

#Fluff 10 IP Range
f10_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f11", shell=True).decode().strip('\n')
f10_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f27", shell=True).decode().strip('\n')
f10_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f27", shell=True).decode().strip('\n')
f10_ip=str(f10_octet1)+'.'+str(f10_octet2)+'.'+str(f10_octet3)+'.0/24'
print(f10_ip)

#Fluff 11 IP Range
f11_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f12", shell=True).decode().strip('\n')
f11_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f28", shell=True).decode().strip('\n')
f11_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f28", shell=True).decode().strip('\n')
f11_ip=str(f11_octet1)+'.'+str(f11_octet2)+'.'+str(f11_octet3)+'.0/24'
print(f11_ip)

#Fluff 12 IP Range
f12_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f13", shell=True).decode().strip('\n')
f12_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f29", shell=True).decode().strip('\n')
f12_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f29", shell=True).decode().strip('\n')
f12_ip=str(f12_octet1)+'.'+str(f12_octet2)+'.'+str(f12_octet3)+'.0/24'
print(f12_ip)

#Fluff 13 IP Range
f13_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f14", shell=True).decode().strip('\n')
f13_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f30", shell=True).decode().strip('\n')
f13_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f30", shell=True).decode().strip('\n')
f13_ip=str(f13_octet1)+'.'+str(f13_octet2)+'.'+str(f13_octet3)+'.0/24'
print(f13_ip)

#Fluff 14 IP Range
f14_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f1", shell=True).decode().strip('\n')
f14_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f31", shell=True).decode().strip('\n')
f14_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f31", shell=True).decode().strip('\n')
f14_ip=str(f14_octet1)+'.'+str(f14_octet2)+'.'+str(f14_octet3)+'.0/24'
print(f14_ip)

#Fluff 15 IP Range
f15_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f2", shell=True).decode().strip('\n')
f15_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f32", shell=True).decode().strip('\n')
f15_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f32", shell=True).decode().strip('\n')
f15_ip=str(f15_octet1)+'.'+str(f15_octet2)+'.'+str(f15_octet3)+'.0/24'
print(f15_ip)

#Fluff 16 IP Range
f16_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f3", shell=True).decode().strip('\n')
f16_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f33", shell=True).decode().strip('\n')
f16_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f33", shell=True).decode().strip('\n')
f16_ip=str(f16_octet1)+'.'+str(f16_octet2)+'.'+str(f16_octet3)+'.0/24'
print(f16_ip)

#Fluff 17 IP Range
f17_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f4", shell=True).decode().strip('\n')
f17_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f34", shell=True).decode().strip('\n')
f17_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f34", shell=True).decode().strip('\n')
f17_ip=str(f17_octet1)+'.'+str(f17_octet2)+'.'+str(f17_octet3)+'.0/24'
print(f17_ip)

#Fluff 18 IP Range
f18_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f5", shell=True).decode().strip('\n')
f18_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f35", shell=True).decode().strip('\n')
f18_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f35", shell=True).decode().strip('\n')
f18_ip=str(f18_octet1)+'.'+str(f18_octet2)+'.'+str(f18_octet3)+'.0/24'
print(f18_ip)

#Fluff 19 IP Range
f19_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f6", shell=True).decode().strip('\n')
f19_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f36", shell=True).decode().strip('\n')
f19_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f36", shell=True).decode().strip('\n')
f19_ip=str(f19_octet1)+'.'+str(f19_octet2)+'.'+str(f19_octet3)+'.0/24'
print(f19_ip)

#Fluff 20 IP Range
f20_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f7", shell=True).decode().strip('\n')
f20_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f37", shell=True).decode().strip('\n')
f20_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f37", shell=True).decode().strip('\n')
f20_ip=str(f20_octet1)+'.'+str(f20_octet2)+'.'+str(f20_octet3)+'.0/24'
print(f20_ip)

#Fluff 21 IP Range
f21_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f8", shell=True).decode().strip('\n')
f21_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f38", shell=True).decode().strip('\n')
f21_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f38", shell=True).decode().strip('\n')
f21_ip=str(f21_octet1)+'.'+str(f21_octet2)+'.'+str(f21_octet3)+'.0/24'
print(f21_ip)

#Fluff 22 IP Range
f22_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f9", shell=True).decode().strip('\n')
f22_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f39", shell=True).decode().strip('\n')
f22_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f39", shell=True).decode().strip('\n')
f22_ip=str(f22_octet1)+'.'+str(f22_octet2)+'.'+str(f22_octet3)+'.0/24'
print(f22_ip)

#Fluff 23 IP Range
f23_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f10", shell=True).decode().strip('\n')
f23_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f40", shell=True).decode().strip('\n')
f23_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f40", shell=True).decode().strip('\n')
f23_ip=str(f23_octet1)+'.'+str(f23_octet2)+'.'+str(f23_octet3)+'.0/24'
print(f23_ip)

#Fluff 24 IP Range
f24_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f11", shell=True).decode().strip('\n')
f24_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f41", shell=True).decode().strip('\n')
f24_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f41", shell=True).decode().strip('\n')
f24_ip=str(f24_octet1)+'.'+str(f24_octet2)+'.'+str(f24_octet3)+'.0/24'
print(f24_ip)

#Fluff 25 IP Range
f25_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f12", shell=True).decode().strip('\n')
f25_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f42", shell=True).decode().strip('\n')
f25_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f42", shell=True).decode().strip('\n')
f25_ip=str(f25_octet1)+'.'+str(f25_octet2)+'.'+str(f25_octet3)+'.0/24'
print(f25_ip)

#Fluff 26 IP Range
f26_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f13", shell=True).decode().strip('\n')
f26_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f43", shell=True).decode().strip('\n')
f26_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f43", shell=True).decode().strip('\n')
f26_ip=str(f26_octet1)+'.'+str(f26_octet2)+'.'+str(f26_octet3)+'.0/24'
print(f26_ip)

#Fluff 27 IP Range
f27_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f14", shell=True).decode().strip('\n')
f27_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f44", shell=True).decode().strip('\n')
f27_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f44", shell=True).decode().strip('\n')
f27_ip=str(f27_octet1)+'.'+str(f27_octet2)+'.'+str(f27_octet3)+'.0/24'
print(f27_ip)

#Fluff 28 IP Range
f28_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f1", shell=True).decode().strip('\n')
f28_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f45", shell=True).decode().strip('\n')
f28_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f45", shell=True).decode().strip('\n')
f28_ip=str(f28_octet1)+'.'+str(f28_octet2)+'.'+str(f28_octet3)+'.0/24'
print(f28_ip)

#Fluff 29 IP Range
f29_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f2", shell=True).decode().strip('\n')
f29_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f46", shell=True).decode().strip('\n')
f29_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f46", shell=True).decode().strip('\n')
f29_ip=str(f29_octet1)+'.'+str(f29_octet2)+'.'+str(f29_octet3)+'.0/24'
print(f29_ip)

#Fluff 30 IP Range
f30_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f3", shell=True).decode().strip('\n')
f30_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f47", shell=True).decode().strip('\n')
f30_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f47", shell=True).decode().strip('\n')
f30_ip=str(f30_octet1)+'.'+str(f30_octet2)+'.'+str(f30_octet3)+'.0/24'
print(f30_ip)

#Fluff 31 IP Range
f31_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f4", shell=True).decode().strip('\n')
f31_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f48", shell=True).decode().strip('\n')
f31_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f48", shell=True).decode().strip('\n')
f31_ip=str(f31_octet1)+'.'+str(f31_octet2)+'.'+str(f31_octet3)+'.0/24'
print(f31_ip)

#Fluff 32 IP Range
f32_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f5", shell=True).decode().strip('\n')
f32_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f49", shell=True).decode().strip('\n')
f32_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f49", shell=True).decode().strip('\n')
f32_ip=str(f32_octet1)+'.'+str(f32_octet2)+'.'+str(f32_octet3)+'.0/24'
print(f32_ip)

#Fluff 33 IP Range
f33_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f6", shell=True).decode().strip('\n')
f33_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f50", shell=True).decode().strip('\n')
f33_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f50", shell=True).decode().strip('\n')
f33_ip=str(f33_octet1)+'.'+str(f33_octet2)+'.'+str(f33_octet3)+'.0/24'
print(f33_ip)

#Fluff 34 IP Range
f34_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f7", shell=True).decode().strip('\n')
f34_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f51", shell=True).decode().strip('\n')
f34_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f51", shell=True).decode().strip('\n')
f34_ip=str(f34_octet1)+'.'+str(f34_octet2)+'.'+str(f34_octet3)+'.0/24'
print(f34_ip)

#Fluff 35 IP Range
f35_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f8", shell=True).decode().strip('\n')
f35_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f52", shell=True).decode().strip('\n')
f35_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f52", shell=True).decode().strip('\n')
f35_ip=str(f35_octet1)+'.'+str(f35_octet2)+'.'+str(f35_octet3)+'.0/24'
print(f35_ip)

#Fluff 36 IP Range
f36_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f9", shell=True).decode().strip('\n')
f36_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f53", shell=True).decode().strip('\n')
f36_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f53", shell=True).decode().strip('\n')
f36_ip=str(f36_octet1)+'.'+str(f36_octet2)+'.'+str(f36_octet3)+'.0/24'
print(f36_ip)

#Fluff 37 IP Range
f37_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f10", shell=True).decode().strip('\n')
f37_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f54", shell=True).decode().strip('\n')
f37_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f54", shell=True).decode().strip('\n')
f37_ip=str(f37_octet1)+'.'+str(f37_octet2)+'.'+str(f37_octet3)+'.0/24'
print(f37_ip)

#Fluff 38 IP Range
f38_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f11", shell=True).decode().strip('\n')
f38_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f55", shell=True).decode().strip('\n')
f38_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f55", shell=True).decode().strip('\n')
f38_ip=str(f38_octet1)+'.'+str(f38_octet2)+'.'+str(f38_octet3)+'.0/24'
print(f38_ip)

#Fluff 39 IP Range
f39_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f12", shell=True).decode().strip('\n')
f39_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f56", shell=True).decode().strip('\n')
f39_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f56", shell=True).decode().strip('\n')
f39_ip=str(f39_octet1)+'.'+str(f39_octet2)+'.'+str(f39_octet3)+'.0/24'
print(f39_ip)

#Fluff 40 IP Range
f40_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f13", shell=True).decode().strip('\n')
f40_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f57", shell=True).decode().strip('\n')
f40_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f57", shell=True).decode().strip('\n')
f40_ip=str(f40_octet1)+'.'+str(f40_octet2)+'.'+str(f40_octet3)+'.0/24'
print(f40_ip)

#Fluff 41 IP Range
f41_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f14", shell=True).decode().strip('\n')
f41_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f58", shell=True).decode().strip('\n')
f41_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f58", shell=True).decode().strip('\n')
f41_ip=str(f41_octet1)+'.'+str(f41_octet2)+'.'+str(f41_octet3)+'.0/24'
print(f41_ip)

#Fluff 42 IP Range
f42_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f1", shell=True).decode().strip('\n')
f42_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f59", shell=True).decode().strip('\n')
f42_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f59", shell=True).decode().strip('\n')
f42_ip=str(f42_octet1)+'.'+str(f42_octet2)+'.'+str(f42_octet3)+'.0/24'
print(f42_ip)

#Fluff 43 IP Range
f43_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f2", shell=True).decode().strip('\n')
f43_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f60", shell=True).decode().strip('\n')
f43_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f60", shell=True).decode().strip('\n')
f43_ip=str(f43_octet1)+'.'+str(f43_octet2)+'.'+str(f43_octet3)+'.0/24'
print(f43_ip)

#Fluff 44 IP Range
f44_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f3", shell=True).decode().strip('\n')
f44_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f61", shell=True).decode().strip('\n')
f44_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f61", shell=True).decode().strip('\n')
f44_ip=str(f44_octet1)+'.'+str(f44_octet2)+'.'+str(f44_octet3)+'.0/24'
print(f44_ip)

#Fluff 45 IP Range
f45_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f4", shell=True).decode().strip('\n')
f45_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f62", shell=True).decode().strip('\n')
f45_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f62", shell=True).decode().strip('\n')
f45_ip=str(f45_octet1)+'.'+str(f45_octet2)+'.'+str(f45_octet3)+'.0/24'
print(f45_ip)

#Fluff 46 IP Range
f46_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f5", shell=True).decode().strip('\n')
f46_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f63", shell=True).decode().strip('\n')
f46_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f63", shell=True).decode().strip('\n')
f46_ip=str(f46_octet1)+'.'+str(f46_octet2)+'.'+str(f46_octet3)+'.0/24'
print(f46_ip)

#Fluff 47 IP Range
f47_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f6", shell=True).decode().strip('\n')
f47_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f64", shell=True).decode().strip('\n')
f47_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f64", shell=True).decode().strip('\n')
f47_ip=str(f47_octet1)+'.'+str(f47_octet2)+'.'+str(f47_octet3)+'.0/24'
print(f47_ip)

#Fluff 48 IP Range
f48_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f7", shell=True).decode().strip('\n')
f48_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f65", shell=True).decode().strip('\n')
f48_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f65", shell=True).decode().strip('\n')
f48_ip=str(f48_octet1)+'.'+str(f48_octet2)+'.'+str(f48_octet3)+'.0/24'
print(f48_ip)

#Fluff 49 IP Range
f49_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f8", shell=True).decode().strip('\n')
f49_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f66", shell=True).decode().strip('\n')
f49_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f66", shell=True).decode().strip('\n')
f49_ip=str(f49_octet1)+'.'+str(f49_octet2)+'.'+str(f49_octet3)+'.0/24'
print(f49_ip)

#Fluff 50 IP Range
f50_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f9", shell=True).decode().strip('\n')
f50_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f67", shell=True).decode().strip('\n')
f50_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f67", shell=True).decode().strip('\n')
f50_ip=str(f50_octet1)+'.'+str(f50_octet2)+'.'+str(f50_octet3)+'.0/24'
print(f50_ip)

#Fluff 51 IP Range
f51_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f10", shell=True).decode().strip('\n')
f51_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f68", shell=True).decode().strip('\n')
f51_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f68", shell=True).decode().strip('\n')
f51_ip=str(f51_octet1)+'.'+str(f51_octet2)+'.'+str(f51_octet3)+'.0/24'
print(f51_ip)

#Fluff 52 IP Range
f52_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f11", shell=True).decode().strip('\n')
f52_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f69", shell=True).decode().strip('\n')
f52_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f69", shell=True).decode().strip('\n')
f52_ip=str(f52_octet1)+'.'+str(f52_octet2)+'.'+str(f52_octet3)+'.0/24'
print(f52_ip)

#Fluff 53 IP Range
f53_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f12", shell=True).decode().strip('\n')
f53_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f70", shell=True).decode().strip('\n')
f53_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f70", shell=True).decode().strip('\n')
f53_ip=str(f53_octet1)+'.'+str(f53_octet2)+'.'+str(f53_octet3)+'.0/24'
print(f53_ip)

#Fluff 54 IP Range
f54_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f13", shell=True).decode().strip('\n')
f54_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f71", shell=True).decode().strip('\n')
f54_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f71", shell=True).decode().strip('\n')
f54_ip=str(f54_octet1)+'.'+str(f54_octet2)+'.'+str(f54_octet3)+'.0/24'
print(f54_ip)

#Fluff 55 IP Range
f55_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f14", shell=True).decode().strip('\n')
f55_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f72", shell=True).decode().strip('\n')
f55_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f72", shell=True).decode().strip('\n')
f55_ip=str(f55_octet1)+'.'+str(f55_octet2)+'.'+str(f55_octet3)+'.0/24'
print(f55_ip)

#Fluff 56 IP Range
f56_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f1", shell=True).decode().strip('\n')
f56_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f73", shell=True).decode().strip('\n')
f56_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f73", shell=True).decode().strip('\n')
f56_ip=str(f56_octet1)+'.'+str(f56_octet2)+'.'+str(f56_octet3)+'.0/24'
print(f56_ip)

#Fluff 57 IP Range
f57_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f2", shell=True).decode().strip('\n')
f57_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f74", shell=True).decode().strip('\n')
f57_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f74", shell=True).decode().strip('\n')
f57_ip=str(f57_octet1)+'.'+str(f57_octet2)+'.'+str(f57_octet3)+'.0/24'
print(f57_ip)

#Fluff 58 IP Range
f58_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f3", shell=True).decode().strip('\n')
f58_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f75", shell=True).decode().strip('\n')
f58_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f75", shell=True).decode().strip('\n')
f58_ip=str(f58_octet1)+'.'+str(f58_octet2)+'.'+str(f58_octet3)+'.0/24'
print(f58_ip)

#Fluff 59 IP Range
f59_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f4", shell=True).decode().strip('\n')
f59_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f76", shell=True).decode().strip('\n')
f59_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f76", shell=True).decode().strip('\n')
f59_ip=str(f59_octet1)+'.'+str(f59_octet2)+'.'+str(f59_octet3)+'.0/24'
print(f59_ip)

#Fluff 60 IP Range
f60_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f5", shell=True).decode().strip('\n')
f60_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f77", shell=True).decode().strip('\n')
f60_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f77", shell=True).decode().strip('\n')
f60_ip=str(f60_octet1)+'.'+str(f60_octet2)+'.'+str(f60_octet3)+'.0/24'
print(f60_ip)

#Fluff 61 IP Range
f61_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f6", shell=True).decode().strip('\n')
f61_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f78", shell=True).decode().strip('\n')
f61_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f78", shell=True).decode().strip('\n')
f61_ip=str(f61_octet1)+'.'+str(f61_octet2)+'.'+str(f61_octet3)+'.0/24'
print(f61_ip)

#Fluff 62 IP Range
f62_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f7", shell=True).decode().strip('\n')
f62_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f79", shell=True).decode().strip('\n')
f62_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f79", shell=True).decode().strip('\n')
f62_ip=str(f62_octet1)+'.'+str(f62_octet2)+'.'+str(f62_octet3)+'.0/24'
print(f62_ip)

#Fluff 63 IP Range
f63_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f8", shell=True).decode().strip('\n')
f63_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f80", shell=True).decode().strip('\n')
f63_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f80", shell=True).decode().strip('\n')
f63_ip=str(f63_octet1)+'.'+str(f63_octet2)+'.'+str(f63_octet3)+'.0/24'
print(f63_ip)

#Fluff 64 IP Range
f64_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f9", shell=True).decode().strip('\n')
f64_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f81", shell=True).decode().strip('\n')
f64_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f81", shell=True).decode().strip('\n')
f64_ip=str(f64_octet1)+'.'+str(f64_octet2)+'.'+str(f64_octet3)+'.0/24'
print(f64_ip)

#Fluff 65 IP Range
f65_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f10", shell=True).decode().strip('\n')
f65_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f82", shell=True).decode().strip('\n')
f65_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f82", shell=True).decode().strip('\n')
f65_ip=str(f65_octet1)+'.'+str(f65_octet2)+'.'+str(f65_octet3)+'.0/24'
print(f65_ip)

#Fluff 66 IP Range
f66_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f11", shell=True).decode().strip('\n')
f66_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f83", shell=True).decode().strip('\n')
f66_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f83", shell=True).decode().strip('\n')
f66_ip=str(f66_octet1)+'.'+str(f66_octet2)+'.'+str(f66_octet3)+'.0/24'
print(f66_ip)

#Fluff 67 IP Range
f67_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f12", shell=True).decode().strip('\n')
f67_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f84", shell=True).decode().strip('\n')
f67_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f84", shell=True).decode().strip('\n')
f67_ip=str(f67_octet1)+'.'+str(f67_octet2)+'.'+str(f67_octet3)+'.0/24'
print(f67_ip)

#Fluff 68 IP Range
f68_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f13", shell=True).decode().strip('\n')
f68_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f85", shell=True).decode().strip('\n')
f68_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f85", shell=True).decode().strip('\n')
f68_ip=str(f68_octet1)+'.'+str(f68_octet2)+'.'+str(f68_octet3)+'.0/24'
print(f68_ip)

#Fluff 69 IP Range
f69_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f14", shell=True).decode().strip('\n')
f69_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f86", shell=True).decode().strip('\n')
f69_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f86", shell=True).decode().strip('\n')
f69_ip=str(f69_octet1)+'.'+str(f69_octet2)+'.'+str(f69_octet3)+'.0/24'
print(f69_ip)

#Fluff 70 IP Range
f70_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f1", shell=True).decode().strip('\n')
f70_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f87", shell=True).decode().strip('\n')
f70_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f87", shell=True).decode().strip('\n')
f70_ip=str(f70_octet1)+'.'+str(f70_octet2)+'.'+str(f70_octet3)+'.0/24'
print(f70_ip)

#Fluff 71 IP Range
f71_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f2", shell=True).decode().strip('\n')
f71_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f88", shell=True).decode().strip('\n')
f71_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f88", shell=True).decode().strip('\n')
f71_ip=str(f71_octet1)+'.'+str(f71_octet2)+'.'+str(f71_octet3)+'.0/24'
print(f71_ip)

#Fluff 72 IP Range
f72_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f3", shell=True).decode().strip('\n')
f72_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f89", shell=True).decode().strip('\n')
f72_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f89", shell=True).decode().strip('\n')
f72_ip=str(f72_octet1)+'.'+str(f72_octet2)+'.'+str(f72_octet3)+'.0/24'
print(f72_ip)

#Fluff 73 IP Range
f73_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f4", shell=True).decode().strip('\n')
f73_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f90", shell=True).decode().strip('\n')
f73_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f90", shell=True).decode().strip('\n')
f73_ip=str(f73_octet1)+'.'+str(f73_octet2)+'.'+str(f73_octet3)+'.0/24'
print(f73_ip)

#Fluff 74 IP Range
f74_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f5", shell=True).decode().strip('\n')
f74_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f91", shell=True).decode().strip('\n')
f74_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f91", shell=True).decode().strip('\n')
f74_ip=str(f74_octet1)+'.'+str(f74_octet2)+'.'+str(f74_octet3)+'.0/24'
print(f74_ip)

#Fluff 75 IP Range
f75_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f6", shell=True).decode().strip('\n')
f75_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f92", shell=True).decode().strip('\n')
f75_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f92", shell=True).decode().strip('\n')
f75_ip=str(f75_octet1)+'.'+str(f75_octet2)+'.'+str(f75_octet3)+'.0/24'
print(f75_ip)

#Fluff 76 IP Range
f76_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f7", shell=True).decode().strip('\n')
f76_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f93", shell=True).decode().strip('\n')
f76_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f93", shell=True).decode().strip('\n')
f76_ip=str(f76_octet1)+'.'+str(f76_octet2)+'.'+str(f76_octet3)+'.0/24'
print(f76_ip)

#Fluff 77 IP Range
f77_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f8", shell=True).decode().strip('\n')
f77_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f94", shell=True).decode().strip('\n')
f77_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f94", shell=True).decode().strip('\n')
f77_ip=str(f77_octet1)+'.'+str(f77_octet2)+'.'+str(f77_octet3)+'.0/24'
print(f77_ip)

#Fluff 78 IP Range
f78_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f9", shell=True).decode().strip('\n')
f78_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f95", shell=True).decode().strip('\n')
f78_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f95", shell=True).decode().strip('\n')
f78_ip=str(f78_octet1)+'.'+str(f78_octet2)+'.'+str(f78_octet3)+'.0/24'
print(f78_ip)

#Fluff 79 IP Range
f79_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f10", shell=True).decode().strip('\n')
f79_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f96", shell=True).decode().strip('\n')
f79_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f96", shell=True).decode().strip('\n')
f79_ip=str(f79_octet1)+'.'+str(f79_octet2)+'.'+str(f79_octet3)+'.0/24'
print(f79_ip)

#Fluff 80 IP Range
f80_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f11", shell=True).decode().strip('\n')
f80_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f97", shell=True).decode().strip('\n')
f80_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f97", shell=True).decode().strip('\n')
f80_ip=str(f80_octet1)+'.'+str(f80_octet2)+'.'+str(f80_octet3)+'.0/24'
print(f80_ip)

#Fluff 81 IP Range
f81_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f12", shell=True).decode().strip('\n')
f81_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f98", shell=True).decode().strip('\n')
f81_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f98", shell=True).decode().strip('\n')
f81_ip=str(f81_octet1)+'.'+str(f81_octet2)+'.'+str(f81_octet3)+'.0/24'
print(f81_ip)

#Fluff 82 IP Range
f82_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f13", shell=True).decode().strip('\n')
f82_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f99", shell=True).decode().strip('\n')
f82_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f99", shell=True).decode().strip('\n')
f82_ip=str(f82_octet1)+'.'+str(f82_octet2)+'.'+str(f82_octet3)+'.0/24'
print(f82_ip)

#Fluff 83 IP Range
f83_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f14", shell=True).decode().strip('\n')
f83_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f100", shell=True).decode().strip('\n')
f83_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f100", shell=True).decode().strip('\n')
f83_ip=str(f83_octet1)+'.'+str(f83_octet2)+'.'+str(f83_octet3)+'.0/24'
print(f83_ip)

#Fluff 84 IP Range
f84_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f1", shell=True).decode().strip('\n')
f84_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f101", shell=True).decode().strip('\n')
f84_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f101", shell=True).decode().strip('\n')
f84_ip=str(f84_octet1)+'.'+str(f84_octet2)+'.'+str(f84_octet3)+'.0/24'
print(f84_ip)

#Fluff 85 IP Range
f85_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f2", shell=True).decode().strip('\n')
f85_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f102", shell=True).decode().strip('\n')
f85_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f102", shell=True).decode().strip('\n')
f85_ip=str(f85_octet1)+'.'+str(f85_octet2)+'.'+str(f85_octet3)+'.0/24'
print(f85_ip)

#Fluff 86 IP Range
f86_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f3", shell=True).decode().strip('\n')
f86_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f103", shell=True).decode().strip('\n')
f86_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f103", shell=True).decode().strip('\n')
f86_ip=str(f86_octet1)+'.'+str(f86_octet2)+'.'+str(f86_octet3)+'.0/24'
print(f86_ip)

#Fluff 87 IP Range
f87_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f4", shell=True).decode().strip('\n')
f87_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f104", shell=True).decode().strip('\n')
f87_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f104", shell=True).decode().strip('\n')
f87_ip=str(f87_octet1)+'.'+str(f87_octet2)+'.'+str(f87_octet3)+'.0/24'
print(f87_ip)

#Fluff 88 IP Range
f88_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f5", shell=True).decode().strip('\n')
f88_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f105", shell=True).decode().strip('\n')
f88_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f105", shell=True).decode().strip('\n')
f88_ip=str(f88_octet1)+'.'+str(f88_octet2)+'.'+str(f88_octet3)+'.0/24'
print(f88_ip)

#Fluff 89 IP Range
f89_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f6", shell=True).decode().strip('\n')
f89_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f106", shell=True).decode().strip('\n')
f89_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f106", shell=True).decode().strip('\n')
f89_ip=str(f89_octet1)+'.'+str(f89_octet2)+'.'+str(f89_octet3)+'.0/24'
print(f89_ip)

#Fluff 90 IP Range
f90_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f7", shell=True).decode().strip('\n')
f90_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f107", shell=True).decode().strip('\n')
f90_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f107", shell=True).decode().strip('\n')
f90_ip=str(f90_octet1)+'.'+str(f90_octet2)+'.'+str(f90_octet3)+'.0/24'
print(f90_ip)

#Fluff 91 IP Range
f91_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f8", shell=True).decode().strip('\n')
f91_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f108", shell=True).decode().strip('\n')
f91_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f108", shell=True).decode().strip('\n')
f91_ip=str(f91_octet1)+'.'+str(f91_octet2)+'.'+str(f91_octet3)+'.0/24'
print(f91_ip)

#Fluff 92 IP Range
f92_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f9", shell=True).decode().strip('\n')
f92_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f109", shell=True).decode().strip('\n')
f92_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f109", shell=True).decode().strip('\n')
f92_ip=str(f92_octet1)+'.'+str(f92_octet2)+'.'+str(f92_octet3)+'.0/24'
print(f92_ip)

#Fluff 93 IP Range
f93_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f10", shell=True).decode().strip('\n')
f93_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f110", shell=True).decode().strip('\n')
f93_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f110", shell=True).decode().strip('\n')
f93_ip=str(f93_octet1)+'.'+str(f93_octet2)+'.'+str(f93_octet3)+'.0/24'
print(f93_ip)

#Fluff 94 IP Range
f94_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f11", shell=True).decode().strip('\n')
f94_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f111", shell=True).decode().strip('\n')
f94_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f111", shell=True).decode().strip('\n')
f94_ip=str(f94_octet1)+'.'+str(f94_octet2)+'.'+str(f94_octet3)+'.0/24'
print(f94_ip)

#Fluff 95 IP Range
f95_octet1=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o1' | cut -d ' ' -f1", shell=True).decode().strip('\n')
f95_octet2=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o2' | cut -d ' ' -f112", shell=True).decode().strip('\n')
f95_octet3=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.o3' | cut -d ' ' -f112", shell=True).decode().strip('\n')
f95_ip=str(f95_octet1)+'.'+str(f95_octet2)+'.'+str(f95_octet3)+'.0/24'
print(f95_ip)

def write_apt(name, date, goal1, goal2, ip, target):
    file=name+'.docx'
    doc = docx.Document()
    doc.add_heading(name, 0)
    para = doc.add_paragraph().add_run(
            'APT Name: '+name)
    para = doc.add_paragraph().add_run(
            'Date First Seen: '+date)
    para = doc.add_paragraph().add_run(
            'Suspected Goals: '+goal1+', '+goal2)
    para = doc.add_paragraph().add_run(
            'Suspected Attacker IPs: '+ip)
    para = doc.add_paragraph().add_run(
            'Suspected Achilles Heel: '+target)
    para.font.size = Pt(12)
    doc.save(file)
    output = subprocess.check_output(['libreoffice', '--convert-to', 'pdf', file])

write_apt(name=c1_apt1, date=c1_apt1_date, goal1=c1_apt1_goal1, goal2=c1_apt1_goal2, ip=a1_ip, target=c1_apt1_target)
write_apt(name=c1_apt2, date=c1_apt2_date, goal1=c1_apt2_goal1, goal2=c1_apt2_goal2, ip=f1_ip, target=c1_apt2_target)
write_apt(name=c1_apt3, date=c1_apt3_date, goal1=c1_apt3_goal1, goal2=c1_apt3_goal2, ip=f2_ip, target=c1_apt3_target)
write_apt(name=c1_apt4, date=c1_apt4_date, goal1=c1_apt4_goal1, goal2=c1_apt4_goal2, ip=f3_ip, target=c1_apt4_target)
write_apt(name=c1_apt5, date=c1_apt5_date, goal1=c1_apt5_goal1, goal2=c1_apt5_goal2, ip=f4_ip, target=c1_apt5_target)
write_apt(name=c1_apt6, date=c1_apt6_date, goal1=c1_apt6_goal1, goal2=c1_apt6_goal2, ip=f5_ip, target=c1_apt6_target)
write_apt(name=c1_apt7, date=c1_apt7_date, goal1=c1_apt7_goal1, goal2=c1_apt7_goal2, ip=f6_ip, target=c1_apt7_target)

write_apt(name=c2_apt1, date=c2_apt1_date, goal1=c2_apt1_goal1, goal2=c2_apt1_goal2, ip=a2_ip, target=c2_apt1_target)
write_apt(name=c2_apt2, date=c2_apt2_date, goal1=c2_apt2_goal1, goal2=c2_apt2_goal2, ip=f7_ip, target=c2_apt2_target)
write_apt(name=c2_apt3, date=c2_apt3_date, goal1=c2_apt3_goal1, goal2=c2_apt3_goal2, ip=f8_ip, target=c2_apt3_target)
write_apt(name=c2_apt4, date=c2_apt4_date, goal1=c2_apt4_goal1, goal2=c2_apt4_goal2, ip=f9_ip, target=c2_apt4_target)
write_apt(name=c2_apt5, date=c2_apt5_date, goal1=c2_apt5_goal1, goal2=c2_apt5_goal2, ip=f10_ip, target=c2_apt5_target)
write_apt(name=c2_apt6, date=c2_apt6_date, goal1=c2_apt6_goal1, goal2=c2_apt6_goal2, ip=f11_ip,target=c2_apt6_target)
write_apt(name=c2_apt7, date=c2_apt7_date, goal1=c2_apt7_goal1, goal2=c2_apt7_goal2, ip=f12_ip, target=c2_apt7_target)

write_apt(name=c3_apt1, date=c3_apt1_date, goal1=c3_apt1_goal1, goal2=c3_apt1_goal2, ip=a3_ip, target=c3_apt1_target)
write_apt(name=c3_apt2, date=c3_apt2_date, goal1=c3_apt2_goal1, goal2=c3_apt2_goal2, ip=f13_ip, target=c3_apt2_target)
write_apt(name=c3_apt3, date=c3_apt3_date, goal1=c3_apt3_goal1, goal2=c3_apt3_goal2, ip=f14_ip, target=c3_apt3_target)
write_apt(name=c3_apt4, date=c3_apt4_date, goal1=c3_apt4_goal1, goal2=c3_apt4_goal2, ip=f15_ip, target=c3_apt4_target)
write_apt(name=c3_apt5, date=c3_apt5_date, goal1=c3_apt5_goal1, goal2=c3_apt5_goal2, ip=f16_ip, target=c3_apt5_target)
write_apt(name=c3_apt6, date=c3_apt6_date, goal1=c3_apt6_goal1, goal2=c3_apt6_goal2, ip=f17_ip, target=c3_apt6_target)
write_apt(name=c3_apt7, date=c3_apt7_date, goal1=c3_apt7_goal1, goal2=c3_apt7_goal2, ip=f18_ip, target=c3_apt7_target)

write_apt(name=c4_apt1, date=c4_apt1_date, goal1=c4_apt1_goal1, goal2=c4_apt1_goal2, ip=f19_ip, target=c4_apt1_target)
write_apt(name=c4_apt2, date=c4_apt2_date, goal1=c4_apt2_goal1, goal2=c4_apt2_goal2, ip=f20_ip, target=c4_apt2_target)
write_apt(name=c4_apt3, date=c4_apt3_date, goal1=c4_apt3_goal1, goal2=c4_apt3_goal2, ip=f21_ip, target=c4_apt3_target)
write_apt(name=c4_apt4, date=c4_apt4_date, goal1=c4_apt4_goal1, goal2=c4_apt4_goal2, ip=f22_ip, target=c4_apt4_target)
write_apt(name=c4_apt5, date=c4_apt5_date, goal1=c4_apt5_goal1, goal2=c4_apt5_goal2, ip=f23_ip, target=c4_apt5_target)
write_apt(name=c4_apt6, date=c4_apt6_date, goal1=c4_apt6_goal1, goal2=c4_apt6_goal2, ip=f24_ip, target=c4_apt6_target)
write_apt(name=c4_apt7, date=c4_apt7_date, goal1=c4_apt7_goal1, goal2=c4_apt7_goal2, ip=f25_ip, target=c4_apt7_target)

write_apt(name=c5_apt1, date=c5_apt1_date, goal1=c5_apt1_goal1, goal2=c5_apt1_goal2, ip=f26_ip, target=c5_apt1_target)
write_apt(name=c5_apt2, date=c5_apt2_date, goal1=c5_apt2_goal1, goal2=c5_apt2_goal2, ip=f27_ip, target=c5_apt2_target)
write_apt(name=c5_apt3, date=c5_apt3_date, goal1=c5_apt3_goal1, goal2=c5_apt3_goal2, ip=f28_ip, target=c5_apt3_target)
write_apt(name=c5_apt4, date=c5_apt4_date, goal1=c5_apt4_goal1, goal2=c5_apt4_goal2, ip=f29_ip, target=c5_apt4_target)
write_apt(name=c5_apt5, date=c5_apt5_date, goal1=c5_apt5_goal1, goal2=c5_apt5_goal2, ip=f30_ip, target=c5_apt5_target)
write_apt(name=c5_apt6, date=c5_apt6_date, goal1=c5_apt6_goal1, goal2=c5_apt6_goal2, ip=f31_ip, target=c5_apt6_target)
write_apt(name=c5_apt7, date=c5_apt7_date, goal1=c5_apt7_goal1, goal2=c5_apt7_goal2, ip=f32_ip, target=c5_apt7_target)

write_apt(name=c6_apt1, date=c6_apt1_date, goal1=c6_apt1_goal1, goal2=c6_apt1_goal2, ip=f33_ip, target=c6_apt1_target)
write_apt(name=c6_apt2, date=c6_apt2_date, goal1=c6_apt2_goal1, goal2=c6_apt2_goal2, ip=f34_ip, target=c6_apt2_target)
write_apt(name=c6_apt3, date=c6_apt3_date, goal1=c6_apt3_goal1, goal2=c6_apt3_goal2, ip=f35_ip, target=c6_apt3_target)
write_apt(name=c6_apt4, date=c6_apt4_date, goal1=c6_apt4_goal1, goal2=c6_apt4_goal2, ip=f36_ip, target=c6_apt4_target)
write_apt(name=c6_apt5, date=c6_apt5_date, goal1=c6_apt5_goal1, goal2=c6_apt5_goal2, ip=f37_ip, target=c6_apt5_target)
write_apt(name=c6_apt6, date=c6_apt6_date, goal1=c6_apt6_goal1, goal2=c6_apt6_goal2, ip=f38_ip, target=c6_apt6_target)
write_apt(name=c6_apt7, date=c6_apt7_date, goal1=c6_apt7_goal1, goal2=c6_apt7_goal2, ip=f39_ip, target=c6_apt7_target)

write_apt(name=c7_apt1, date=c7_apt1_date, goal1=c7_apt1_goal1, goal2=c7_apt1_goal2, ip=f40_ip, target=c7_apt1_target)
write_apt(name=c7_apt2, date=c7_apt2_date, goal1=c7_apt2_goal1, goal2=c7_apt2_goal2, ip=f41_ip, target=c7_apt2_target)
write_apt(name=c7_apt3, date=c7_apt3_date, goal1=c7_apt3_goal1, goal2=c7_apt3_goal2, ip=f42_ip, target=c7_apt3_target)
write_apt(name=c7_apt4, date=c7_apt4_date, goal1=c7_apt4_goal1, goal2=c7_apt4_goal2, ip=f43_ip, target=c7_apt4_target)
write_apt(name=c7_apt5, date=c7_apt5_date, goal1=c7_apt5_goal1, goal2=c7_apt5_goal2, ip=f44_ip, target=c7_apt5_target)
write_apt(name=c7_apt6, date=c7_apt6_date, goal1=c7_apt6_goal1, goal2=c7_apt6_goal2, ip=f45_ip, target=c7_apt6_target)
write_apt(name=c7_apt7, date=c7_apt7_date, goal1=c7_apt7_goal1, goal2=c7_apt7_goal2, ip=f46_ip, target=c7_apt7_target)

write_apt(name=c8_apt1, date=c8_apt1_date, goal1=c8_apt1_goal1, goal2=c8_apt1_goal2, ip=f47_ip, target=c8_apt1_target)
write_apt(name=c8_apt2, date=c8_apt2_date, goal1=c8_apt2_goal1, goal2=c8_apt2_goal2, ip=f48_ip, target=c8_apt2_target)
write_apt(name=c8_apt3, date=c8_apt3_date, goal1=c8_apt3_goal1, goal2=c8_apt3_goal2, ip=f49_ip, target=c8_apt3_target)
write_apt(name=c8_apt4, date=c8_apt4_date, goal1=c8_apt4_goal1, goal2=c8_apt4_goal2, ip=f50_ip, target=c8_apt4_target)
write_apt(name=c8_apt5, date=c8_apt5_date, goal1=c8_apt5_goal1, goal2=c8_apt5_goal2, ip=f51_ip, target=c8_apt5_target)
write_apt(name=c8_apt6, date=c8_apt6_date, goal1=c8_apt6_goal1, goal2=c8_apt6_goal2, ip=f52_ip, target=c8_apt6_target)
write_apt(name=c8_apt7, date=c8_apt7_date, goal1=c8_apt7_goal1, goal2=c8_apt7_goal2, ip=f53_ip, target=c8_apt7_target)

write_apt(name=c9_apt1, date=c9_apt1_date, goal1=c9_apt1_goal1, goal2=c9_apt1_goal2, ip=f54_ip, target=c9_apt1_target)
write_apt(name=c9_apt2, date=c9_apt2_date, goal1=c9_apt2_goal1, goal2=c9_apt2_goal2, ip=f55_ip, target=c9_apt2_target)
write_apt(name=c9_apt3, date=c9_apt3_date, goal1=c9_apt3_goal1, goal2=c9_apt3_goal2, ip=f56_ip, target=c9_apt3_target)
write_apt(name=c9_apt4, date=c9_apt4_date, goal1=c9_apt4_goal1, goal2=c9_apt4_goal2, ip=f57_ip, target=c9_apt4_target)
write_apt(name=c9_apt5, date=c9_apt5_date, goal1=c9_apt5_goal1, goal2=c9_apt5_goal2, ip=f58_ip, target=c9_apt5_target)
write_apt(name=c9_apt6, date=c9_apt6_date, goal1=c9_apt6_goal1, goal2=c9_apt6_goal2, ip=f59_ip, target=c9_apt6_target)
write_apt(name=c9_apt7, date=c9_apt7_date, goal1=c9_apt7_goal1, goal2=c9_apt7_goal2, ip=f60_ip, target=c9_apt7_target)

write_apt(name=c10_apt1, date=c10_apt1_date, goal1=c10_apt1_goal1, goal2=c10_apt1_goal2, ip=f61_ip, target=c10_apt1_target)
write_apt(name=c10_apt2, date=c10_apt2_date, goal1=c10_apt2_goal1, goal2=c10_apt2_goal2, ip=f62_ip, target=c10_apt2_target)
write_apt(name=c10_apt3, date=c10_apt3_date, goal1=c10_apt3_goal1, goal2=c10_apt3_goal2, ip=f63_ip, target=c10_apt3_target)
write_apt(name=c10_apt4, date=c10_apt4_date, goal1=c10_apt4_goal1, goal2=c10_apt4_goal2, ip=f64_ip, target=c10_apt4_target)
write_apt(name=c10_apt5, date=c10_apt5_date, goal1=c10_apt5_goal1, goal2=c10_apt5_goal2, ip=f65_ip, target=c10_apt5_target)
write_apt(name=c10_apt6, date=c10_apt6_date, goal1=c10_apt6_goal1, goal2=c10_apt6_goal2, ip=f66_ip, target=c10_apt6_target)
write_apt(name=c10_apt7, date=c10_apt7_date, goal1=c10_apt7_goal1, goal2=c10_apt7_goal2, ip=f67_ip, target=c10_apt7_target)

write_apt(name=c11_apt1, date=c11_apt1_date, goal1=c11_apt1_goal1, goal2=c11_apt1_goal2, ip=f68_ip, target=c11_apt1_target)
write_apt(name=c11_apt2, date=c11_apt2_date, goal1=c11_apt2_goal1, goal2=c11_apt2_goal2, ip=f69_ip, target=c11_apt2_target)
write_apt(name=c11_apt3, date=c11_apt3_date, goal1=c11_apt3_goal1, goal2=c11_apt3_goal2, ip=f70_ip, target=c11_apt3_target)
write_apt(name=c11_apt4, date=c11_apt4_date, goal1=c11_apt4_goal1, goal2=c11_apt4_goal2, ip=f71_ip, target=c11_apt4_target)
write_apt(name=c11_apt5, date=c11_apt5_date, goal1=c11_apt5_goal1, goal2=c11_apt5_goal2, ip=f72_ip, target=c11_apt5_target)
write_apt(name=c11_apt6, date=c11_apt6_date, goal1=c11_apt6_goal1, goal2=c11_apt6_goal2, ip=f73_ip, target=c11_apt6_target)
write_apt(name=c11_apt7, date=c11_apt7_date, goal1=c11_apt7_goal1, goal2=c11_apt7_goal2, ip=f74_ip, target=c11_apt7_target)

write_apt(name=c12_apt1, date=c12_apt1_date, goal1=c12_apt1_goal1, goal2=c12_apt1_goal2, ip=f75_ip, target=c12_apt1_target)
write_apt(name=c12_apt2, date=c12_apt2_date, goal1=c12_apt2_goal1, goal2=c12_apt2_goal2, ip=f76_ip, target=c12_apt2_target)
write_apt(name=c12_apt3, date=c12_apt3_date, goal1=c12_apt3_goal1, goal2=c12_apt3_goal2, ip=f77_ip, target=c12_apt3_target)
write_apt(name=c12_apt4, date=c12_apt4_date, goal1=c12_apt4_goal1, goal2=c12_apt4_goal2, ip=f78_ip, target=c12_apt4_target)
write_apt(name=c12_apt5, date=c12_apt5_date, goal1=c12_apt5_goal1, goal2=c12_apt5_goal2, ip=f79_ip, target=c12_apt5_target)
write_apt(name=c12_apt6, date=c12_apt6_date, goal1=c12_apt6_goal1, goal2=c12_apt6_goal2, ip=f80_ip, target=c12_apt6_target)
write_apt(name=c12_apt7, date=c12_apt7_date, goal1=c12_apt7_goal1, goal2=c12_apt7_goal2, ip=f81_ip, target=c12_apt7_target)

write_apt(name=c13_apt1, date=c13_apt1_date, goal1=c13_apt1_goal1, goal2=c13_apt1_goal2, ip=f82_ip, target=c13_apt1_target)
write_apt(name=c13_apt2, date=c13_apt2_date, goal1=c13_apt2_goal1, goal2=c13_apt2_goal2, ip=f83_ip, target=c13_apt2_target)
write_apt(name=c13_apt3, date=c13_apt3_date, goal1=c13_apt3_goal1, goal2=c13_apt3_goal2, ip=f84_ip, target=c13_apt3_target)
write_apt(name=c13_apt4, date=c13_apt4_date, goal1=c13_apt4_goal1, goal2=c13_apt4_goal2, ip=f85_ip, target=c13_apt4_target)
write_apt(name=c13_apt5, date=c13_apt5_date, goal1=c13_apt5_goal1, goal2=c13_apt5_goal2, ip=f86_ip, target=c13_apt5_target)
write_apt(name=c13_apt6, date=c13_apt6_date, goal1=c13_apt6_goal1, goal2=c13_apt6_goal2, ip=f87_ip, target=c13_apt6_target)
write_apt(name=c13_apt7, date=c13_apt7_date, goal1=c13_apt7_goal1, goal2=c13_apt7_goal2, ip=f88_ip, target=c13_apt7_target)

write_apt(name=c14_apt1, date=c14_apt1_date, goal1=c14_apt1_goal1, goal2=c14_apt1_goal2, ip=f89_ip, target=c14_apt1_target)
write_apt(name=c14_apt2, date=c14_apt2_date, goal1=c14_apt2_goal1, goal2=c14_apt2_goal2, ip=f90_ip, target=c14_apt2_target)
write_apt(name=c14_apt3, date=c14_apt3_date, goal1=c14_apt3_goal1, goal2=c14_apt3_goal2, ip=f91_ip, target=c14_apt3_target)
write_apt(name=c14_apt4, date=c14_apt4_date, goal1=c14_apt4_goal1, goal2=c14_apt4_goal2, ip=f92_ip, target=c14_apt4_target)
write_apt(name=c14_apt5, date=c14_apt5_date, goal1=c14_apt5_goal1, goal2=c14_apt5_goal2, ip=f93_ip, target=c14_apt5_target)
write_apt(name=c14_apt6, date=c14_apt6_date, goal1=c14_apt6_goal1, goal2=c14_apt6_goal2, ip=f94_ip, target=c14_apt6_target)
write_apt(name=c14_apt7, date=c14_apt7_date, goal1=c14_apt7_goal1, goal2=c14_apt7_goal2, ip=f95_ip, target=c14_apt7_target)


# Print Country 1 APT1
print(f'APT Name: {c1_apt1}')
print(f'Goals: {c1_apt1_goal1}, {c1_apt1_goal2}')
print(f'Date First Seen: {c1_apt1_date}')
print(f'Achilles Heel: {c1_apt1_target}')

# Print Country 1 APT2
print(f'APT Name: {c1_apt2}')
print(f'Goals: {c1_apt2_goal1}, {c1_apt2_goal2}')
print(f'Date First Seen: {c1_apt2_date}')
print(f'Achilles Heel: {c1_apt2_target}')

# Print Country 1 APT3
print(f'APT Name: {c1_apt3}')
print(f'Goals: {c1_apt3_goal1}, {c1_apt3_goal2}')
print(f'Date First Seen: {c1_apt3_date}')
print(f'Achilles Heel: {c1_apt3_target}')

# Print Country 1 APT4
print(f'APT Name: {c1_apt4}')
print(f'Goals: {c1_apt4_goal1}, {c1_apt4_goal2}')
print(f'Date First Seen: {c1_apt4_date}')
print(f'Achilles Heel: {c1_apt4_target}')

# Print Country 1 APT5
print(f'APT Name: {c1_apt5}')
print(f'Goals: {c1_apt5_goal1}, {c1_apt5_goal2}')
print(f'Date First Seen: {c1_apt5_date}')
print(f'Achilles Heel: {c1_apt5_target}')

# Print Country 1 APT6
print(f'APT Name: {c1_apt6}')
print(f'Goals: {c1_apt6_goal1}, {c1_apt6_goal2}')
print(f'Date First Seen: {c1_apt6_date}')
print(f'Achilles Heel: {c1_apt6_target}')

# Print Country 1 APT7
print(f'APT Name: {c1_apt7}')
print(f'Goals: {c1_apt7_goal1}, {c1_apt7_goal2}')
print(f'Date First Seen: {c1_apt7_date}')
print(f'Achilles Heel: {c1_apt7_target}')


# Print Country 2 APT1
print(f'APT Name: {c2_apt1}')
print(f'Goals: {c2_apt1_goal1}, {c2_apt1_goal2}')
print(f'Date First Seen: {c2_apt1_date}')
print(f'Achilles Heel: {c2_apt1_target}')

# Print Country 2 APT2
print(f'APT Name: {c2_apt2}')
print(f'Goals: {c2_apt2_goal1}, {c2_apt2_goal2}')
print(f'Date First Seen: {c2_apt2_date}')
print(f'Achilles Heel: {c2_apt2_target}')

# Print Country 2 APT3
print(f'APT Name: {c2_apt3}')
print(f'Goals: {c2_apt3_goal1}, {c2_apt3_goal2}')
print(f'Date First Seen: {c2_apt3_date}')
print(f'Achilles Heel: {c2_apt3_target}')

# Print Country 2 APT4
print(f'APT Name: {c2_apt4}')
print(f'Goals: {c2_apt4_goal1}, {c2_apt4_goal2}')
print(f'Date First Seen: {c2_apt4_date}')
print(f'Achilles Heel: {c2_apt4_target}')

# Print Country 2 APT5
print(f'APT Name: {c2_apt5}')
print(f'Goals: {c2_apt5_goal1}, {c2_apt5_goal2}')
print(f'Date First Seen: {c2_apt5_date}')
print(f'Achilles Heel: {c2_apt5_target}')

# Print Country 2 APT6
print(f'APT Name: {c2_apt6}')
print(f'Goals: {c2_apt6_goal1}, {c2_apt6_goal2}')
print(f'Date First Seen: {c2_apt6_date}')
print(f'Achilles Heel: {c2_apt6_target}')

# Print Country 2 APT7
print(f'APT Name: {c2_apt7}')
print(f'Goals: {c2_apt7_goal1}, {c2_apt7_goal2}')
print(f'Date First Seen: {c2_apt7_date}')
print(f'Achilles Heel: {c2_apt7_target}')


# Print Country 3 APT1
print(f'APT Name: {c3_apt1}')
print(f'Goals: {c3_apt1_goal1}, {c3_apt1_goal2}')
print(f'Date First Seen: {c3_apt1_date}')
print(f'Achilles Heel: {c3_apt1_target}')

# Print Country 3 APT2
print(f'APT Name: {c3_apt2}')
print(f'Goals: {c3_apt2_goal1}, {c3_apt2_goal2}')
print(f'Date First Seen: {c3_apt2_date}')
print(f'Achilles Heel: {c3_apt2_target}')

# Print Country 3 APT3
print(f'APT Name: {c3_apt3}')
print(f'Goals: {c3_apt3_goal1}, {c3_apt3_goal2}')
print(f'Date First Seen: {c3_apt3_date}')
print(f'Achilles Heel: {c3_apt3_target}')

# Print Country 3 APT4
print(f'APT Name: {c3_apt4}')
print(f'Goals: {c3_apt4_goal1}, {c3_apt4_goal2}')
print(f'Date First Seen: {c3_apt4_date}')
print(f'Achilles Heel: {c3_apt4_target}')

# Print Country 3 APT5
print(f'APT Name: {c3_apt5}')
print(f'Goals: {c3_apt5_goal1}, {c3_apt5_goal2}')
print(f'Date First Seen: {c3_apt5_date}')
print(f'Achilles Heel: {c3_apt5_target}')

# Print Country 3 APT6
print(f'APT Name: {c3_apt6}')
print(f'Goals: {c3_apt6_goal1}, {c3_apt6_goal2}')
print(f'Date First Seen: {c3_apt6_date}')
print(f'Achilles Heel: {c3_apt6_target}')

# Print Country 3 APT7
print(f'APT Name: {c3_apt7}')
print(f'Goals: {c3_apt7_goal1}, {c3_apt7_goal2}')
print(f'Date First Seen: {c3_apt7_date}')
print(f'Achilles Heel: {c3_apt7_target}')


# Print Country 4 APT1
print(f'APT Name: {c4_apt1}')
print(f'Goals: {c4_apt1_goal1}, {c4_apt1_goal2}')
print(f'Date First Seen: {c4_apt1_date}')
print(f'Achilles Heel: {c4_apt1_target}')

# Print Country 4 APT2
print(f'APT Name: {c4_apt2}')
print(f'Goals: {c4_apt2_goal1}, {c4_apt2_goal2}')
print(f'Date First Seen: {c4_apt2_date}')
print(f'Achilles Heel: {c4_apt2_target}')

# Print Country 4 APT3
print(f'APT Name: {c4_apt3}')
print(f'Goals: {c4_apt3_goal1}, {c4_apt3_goal2}')
print(f'Date First Seen: {c4_apt3_date}')
print(f'Achilles Heel: {c4_apt3_target}')

# Print Country 4 APT4
print(f'APT Name: {c4_apt4}')
print(f'Goals: {c4_apt4_goal1}, {c4_apt4_goal2}')
print(f'Date First Seen: {c4_apt4_date}')
print(f'Achilles Heel: {c4_apt4_target}')

# Print Country 4 APT5
print(f'APT Name: {c4_apt5}')
print(f'Goals: {c4_apt5_goal1}, {c4_apt5_goal2}')
print(f'Date First Seen: {c4_apt5_date}')
print(f'Achilles Heel: {c4_apt5_target}')

# Print Country 4 APT6
print(f'APT Name: {c4_apt6}')
print(f'Goals: {c4_apt6_goal1}, {c4_apt6_goal2}')
print(f'Date First Seen: {c4_apt6_date}')
print(f'Achilles Heel: {c4_apt6_target}')

# Print Country 4 APT7
print(f'APT Name: {c4_apt7}')
print(f'Goals: {c4_apt7_goal1}, {c4_apt7_goal2}')
print(f'Date First Seen: {c4_apt7_date}')
print(f'Achilles Heel: {c4_apt7_target}')


# Print Country 5 APT1
print(f'APT Name: {c5_apt1}')
print(f'Goals: {c5_apt1_goal1}, {c5_apt1_goal2}')
print(f'Date First Seen: {c5_apt1_date}')
print(f'Achilles Heel: {c5_apt1_target}')

# Print Country 5 APT2
print(f'APT Name: {c5_apt2}')
print(f'Goals: {c5_apt2_goal1}, {c5_apt2_goal2}')
print(f'Date First Seen: {c5_apt2_date}')
print(f'Achilles Heel: {c5_apt2_target}')

# Print Country 5 APT3
print(f'APT Name: {c5_apt3}')
print(f'Goals: {c5_apt3_goal1}, {c5_apt3_goal2}')
print(f'Date First Seen: {c5_apt3_date}')
print(f'Achilles Heel: {c5_apt3_target}')

# Print Country 5 APT4
print(f'APT Name: {c5_apt4}')
print(f'Goals: {c5_apt4_goal1}, {c5_apt4_goal2}')
print(f'Date First Seen: {c5_apt4_date}')
print(f'Achilles Heel: {c5_apt4_target}')

# Print Country 5 APT5
print(f'APT Name: {c5_apt5}')
print(f'Goals: {c5_apt5_goal1}, {c5_apt5_goal2}')
print(f'Date First Seen: {c5_apt5_date}')
print(f'Achilles Heel: {c5_apt5_target}')

# Print Country 5 APT6
print(f'APT Name: {c5_apt6}')
print(f'Goals: {c5_apt6_goal1}, {c5_apt6_goal2}')
print(f'Date First Seen: {c5_apt6_date}')
print(f'Achilles Heel: {c5_apt6_target}')

# Print Country 5 APT7
print(f'APT Name: {c5_apt7}')
print(f'Goals: {c5_apt7_goal1}, {c5_apt7_goal2}')
print(f'Date First Seen: {c5_apt7_date}')
print(f'Achilles Heel: {c5_apt7_target}')


# Print Country 6 APT1
print(f'APT Name: {c6_apt1}')
print(f'Goals: {c6_apt1_goal1}, {c6_apt1_goal2}')
print(f'Date First Seen: {c6_apt1_date}')
print(f'Achilles Heel: {c6_apt1_target}')

# Print Country 6 APT2
print(f'APT Name: {c6_apt2}')
print(f'Goals: {c6_apt2_goal1}, {c6_apt2_goal2}')
print(f'Date First Seen: {c6_apt2_date}')
print(f'Achilles Heel: {c6_apt2_target}')

# Print Country 6 APT3
print(f'APT Name: {c6_apt3}')
print(f'Goals: {c6_apt3_goal1}, {c6_apt3_goal2}')
print(f'Date First Seen: {c6_apt3_date}')
print(f'Achilles Heel: {c6_apt3_target}')

# Print Country 6 APT4
print(f'APT Name: {c6_apt4}')
print(f'Goals: {c6_apt4_goal1}, {c6_apt4_goal2}')
print(f'Date First Seen: {c6_apt4_date}')
print(f'Achilles Heel: {c6_apt4_target}')

# Print Country 6 APT5
print(f'APT Name: {c6_apt5}')
print(f'Goals: {c6_apt5_goal1}, {c6_apt5_goal2}')
print(f'Date First Seen: {c6_apt5_date}')
print(f'Achilles Heel: {c6_apt5_target}')

# Print Country 6 APT6
print(f'APT Name: {c6_apt6}')
print(f'Goals: {c6_apt6_goal1}, {c6_apt6_goal2}')
print(f'Date First Seen: {c6_apt6_date}')
print(f'Achilles Heel: {c6_apt6_target}')

# Print Country 6 APT7
print(f'APT Name: {c6_apt7}')
print(f'Goals: {c6_apt7_goal1}, {c6_apt7_goal2}')
print(f'Date First Seen: {c6_apt7_date}')
print(f'Achilles Heel: {c6_apt7_target}')


# Print Country 7 APT1
print(f'APT Name: {c7_apt1}')
print(f'Goals: {c7_apt1_goal1}, {c7_apt1_goal2}')
print(f'Date First Seen: {c7_apt1_date}')
print(f'Achilles Heel: {c7_apt1_target}')

# Print Country 7 APT2
print(f'APT Name: {c7_apt2}')
print(f'Goals: {c7_apt2_goal1}, {c7_apt2_goal2}')
print(f'Date First Seen: {c7_apt2_date}')
print(f'Achilles Heel: {c7_apt2_target}')

# Print Country 7 APT3
print(f'APT Name: {c7_apt3}')
print(f'Goals: {c7_apt3_goal1}, {c7_apt3_goal2}')
print(f'Date First Seen: {c7_apt3_date}')
print(f'Achilles Heel: {c7_apt3_target}')

# Print Country 7 APT4
print(f'APT Name: {c7_apt4}')
print(f'Goals: {c7_apt4_goal1}, {c7_apt4_goal2}')
print(f'Date First Seen: {c7_apt4_date}')
print(f'Achilles Heel: {c7_apt4_target}')

# Print Country 7 APT5
print(f'APT Name: {c7_apt5}')
print(f'Goals: {c7_apt5_goal1}, {c7_apt5_goal2}')
print(f'Date First Seen: {c7_apt5_date}')
print(f'Achilles Heel: {c7_apt5_target}')

# Print Country 7 APT6
print(f'APT Name: {c7_apt6}')
print(f'Goals: {c7_apt6_goal1}, {c7_apt6_goal2}')
print(f'Date First Seen: {c7_apt6_date}')
print(f'Achilles Heel: {c7_apt6_target}')

# Print Country 7 APT7
print(f'APT Name: {c7_apt7}')
print(f'Goals: {c7_apt7_goal1}, {c7_apt7_goal2}')
print(f'Date First Seen: {c7_apt7_date}')
print(f'Achilles Heel: {c7_apt7_target}')

# Print Country 8 APT1
print(f'APT Name: {c8_apt1}')
print(f'Goals: {c8_apt1_goal1}, {c8_apt1_goal2}')
print(f'Date First Seen: {c8_apt1_date}')
print(f'Achilles Heel: {c8_apt1_target}')

# Print Country 8 APT2
print(f'APT Name: {c8_apt2}')
print(f'Goals: {c8_apt2_goal1}, {c8_apt2_goal2}')
print(f'Date First Seen: {c8_apt2_date}')
print(f'Achilles Heel: {c8_apt2_target}')

# Print Country 8 APT3
print(f'APT Name: {c8_apt3}')
print(f'Goals: {c8_apt3_goal1}, {c8_apt3_goal2}')
print(f'Date First Seen: {c8_apt3_date}')
print(f'Achilles Heel: {c8_apt3_target}')

# Print Country 8 APT4
print(f'APT Name: {c8_apt4}')
print(f'Goals: {c8_apt4_goal1}, {c8_apt4_goal2}')
print(f'Date First Seen: {c8_apt4_date}')
print(f'Achilles Heel: {c8_apt4_target}')

# Print Country 8 APT5
print(f'APT Name: {c8_apt5}')
print(f'Goals: {c8_apt5_goal1}, {c8_apt5_goal2}')
print(f'Date First Seen: {c8_apt5_date}')
print(f'Achilles Heel: {c8_apt5_target}')

# Print Country 8 APT6
print(f'APT Name: {c8_apt6}')
print(f'Goals: {c8_apt6_goal1}, {c8_apt6_goal2}')
print(f'Date First Seen: {c8_apt6_date}')
print(f'Achilles Heel: {c8_apt6_target}')

# Print Country 8 APT7
print(f'APT Name: {c8_apt7}')
print(f'Goals: {c8_apt7_goal1}, {c8_apt7_goal2}')
print(f'Date First Seen: {c8_apt7_date}')
print(f'Achilles Heel: {c8_apt7_target}')


# Print Country 9 APT1
print(f'APT Name: {c9_apt1}')
print(f'Goals: {c9_apt1_goal1}, {c9_apt1_goal2}')
print(f'Date First Seen: {c9_apt1_date}')
print(f'Achilles Heel: {c9_apt1_target}')

# Print Country 9 APT2
print(f'APT Name: {c9_apt2}')
print(f'Goals: {c9_apt2_goal1}, {c9_apt2_goal2}')
print(f'Date First Seen: {c9_apt2_date}')
print(f'Achilles Heel: {c9_apt2_target}')

# Print Country 9 APT3
print(f'APT Name: {c9_apt3}')
print(f'Goals: {c9_apt3_goal1}, {c9_apt3_goal2}')
print(f'Date First Seen: {c9_apt3_date}')
print(f'Achilles Heel: {c9_apt3_target}')

# Print Country 9 APT4
print(f'APT Name: {c9_apt4}')
print(f'Goals: {c9_apt4_goal1}, {c9_apt4_goal2}')
print(f'Date First Seen: {c9_apt4_date}')
print(f'Achilles Heel: {c9_apt4_target}')

# Print Country 9 APT5
print(f'APT Name: {c9_apt5}')
print(f'Goals: {c9_apt5_goal1}, {c9_apt5_goal2}')
print(f'Date First Seen: {c9_apt5_date}')
print(f'Achilles Heel: {c9_apt5_target}')

# Print Country 9 APT6
print(f'APT Name: {c9_apt6}')
print(f'Goals: {c9_apt6_goal1}, {c9_apt6_goal2}')
print(f'Date First Seen: {c9_apt6_date}')
print(f'Achilles Heel: {c9_apt6_target}')

# Print Country 9 APT7
print(f'APT Name: {c9_apt7}')
print(f'Goals: {c9_apt7_goal1}, {c9_apt7_goal2}')
print(f'Date First Seen: {c9_apt7_date}')
print(f'Achilles Heel: {c9_apt7_target}')


# Print Country 10 APT1
print(f'APT Name: {c10_apt1}')
print(f'Goals: {c10_apt1_goal1}, {c10_apt1_goal2}')
print(f'Date First Seen: {c10_apt1_date}')
print(f'Achilles Heel: {c10_apt1_target}')

# Print Country 10 APT2
print(f'APT Name: {c10_apt2}')
print(f'Goals: {c10_apt2_goal1}, {c10_apt2_goal2}')
print(f'Date First Seen: {c10_apt2_date}')
print(f'Achilles Heel: {c10_apt2_target}')

# Print Country 10 APT3
print(f'APT Name: {c10_apt3}')
print(f'Goals: {c10_apt3_goal1}, {c10_apt3_goal2}')
print(f'Date First Seen: {c10_apt3_date}')
print(f'Achilles Heel: {c10_apt3_target}')

# Print Country 10 APT4
print(f'APT Name: {c10_apt4}')
print(f'Goals: {c10_apt4_goal1}, {c10_apt4_goal2}')
print(f'Date First Seen: {c10_apt4_date}')
print(f'Achilles Heel: {c10_apt4_target}')

# Print Country 10 APT5
print(f'APT Name: {c10_apt5}')
print(f'Goals: {c10_apt5_goal1}, {c10_apt5_goal2}')
print(f'Date First Seen: {c10_apt5_date}')
print(f'Achilles Heel: {c10_apt5_target}')

# Print Country 10 APT6
print(f'APT Name: {c10_apt6}')
print(f'Goals: {c10_apt6_goal1}, {c10_apt6_goal2}')
print(f'Date First Seen: {c10_apt6_date}')
print(f'Achilles Heel: {c10_apt6_target}')

# Print Country 10 APT7
print(f'APT Name: {c10_apt7}')
print(f'Goals: {c10_apt7_goal1}, {c10_apt7_goal2}')
print(f'Date First Seen: {c10_apt7_date}')
print(f'Achilles Heel: {c10_apt7_target}')


# Print Country 11 APT1
print(f'APT Name: {c11_apt1}')
print(f'Goals: {c11_apt1_goal1}, {c11_apt1_goal2}')
print(f'Date First Seen: {c11_apt1_date}')
print(f'Achilles Heel: {c11_apt1_target}')

# Print Country 11 APT2
print(f'APT Name: {c11_apt2}')
print(f'Goals: {c11_apt2_goal1}, {c11_apt2_goal2}')
print(f'Date First Seen: {c11_apt2_date}')
print(f'Achilles Heel: {c11_apt2_target}')

# Print Country 11 APT3
print(f'APT Name: {c11_apt3}')
print(f'Goals: {c11_apt3_goal1}, {c11_apt3_goal2}')
print(f'Date First Seen: {c11_apt3_date}')
print(f'Achilles Heel: {c11_apt3_target}')

# Print Country 11 APT4
print(f'APT Name: {c11_apt4}')
print(f'Goals: {c11_apt4_goal1}, {c11_apt4_goal2}')
print(f'Date First Seen: {c11_apt4_date}')
print(f'Achilles Heel: {c11_apt4_target}')

# Print Country 11 APT5
print(f'APT Name: {c11_apt5}')
print(f'Goals: {c11_apt5_goal1}, {c11_apt5_goal2}')
print(f'Date First Seen: {c11_apt5_date}')
print(f'Achilles Heel: {c11_apt5_target}')

# Print Country 11 APT6
print(f'APT Name: {c11_apt6}')
print(f'Goals: {c11_apt6_goal1}, {c11_apt6_goal2}')
print(f'Date First Seen: {c11_apt6_date}')
print(f'Achilles Heel: {c11_apt6_target}')

# Print Country 11 APT7
print(f'APT Name: {c11_apt7}')
print(f'Goals: {c11_apt7_goal1}, {c11_apt7_goal2}')
print(f'Date First Seen: {c11_apt7_date}')
print(f'Achilles Heel: {c11_apt7_target}')


# Print Country 12 APT1
print(f'APT Name: {c12_apt1}')
print(f'Goals: {c12_apt1_goal1}, {c12_apt1_goal2}')
print(f'Date First Seen: {c12_apt1_date}')
print(f'Achilles Heel: {c12_apt1_target}')

# Print Country 12 APT2
print(f'APT Name: {c12_apt2}')
print(f'Goals: {c12_apt2_goal1}, {c12_apt2_goal2}')
print(f'Date First Seen: {c12_apt2_date}')
print(f'Achilles Heel: {c12_apt2_target}')

# Print Country 12 APT3
print(f'APT Name: {c12_apt3}')
print(f'Goals: {c12_apt3_goal1}, {c12_apt3_goal2}')
print(f'Date First Seen: {c12_apt3_date}')
print(f'Achilles Heel: {c12_apt3_target}')

# Print Country 12 APT4
print(f'APT Name: {c12_apt4}')
print(f'Goals: {c12_apt4_goal1}, {c12_apt4_goal2}')
print(f'Date First Seen: {c12_apt4_date}')
print(f'Achilles Heel: {c12_apt4_target}')

# Print Country 12 APT5
print(f'APT Name: {c12_apt5}')
print(f'Goals: {c12_apt5_goal1}, {c12_apt5_goal2}')
print(f'Date First Seen: {c12_apt5_date}')
print(f'Achilles Heel: {c12_apt5_target}')

# Print Country 12 APT6
print(f'APT Name: {c12_apt6}')
print(f'Goals: {c12_apt6_goal1}, {c12_apt6_goal2}')
print(f'Date First Seen: {c12_apt6_date}')
print(f'Achilles Heel: {c12_apt6_target}')

# Print Country 12 APT7
print(f'APT Name: {c12_apt7}')
print(f'Goals: {c12_apt7_goal1}, {c12_apt7_goal2}')
print(f'Date First Seen: {c12_apt7_date}')
print(f'Achilles Heel: {c12_apt7_target}')


# Print Country 13 APT1
print(f'APT Name: {c13_apt1}')
print(f'Goals: {c13_apt1_goal1}, {c13_apt1_goal2}')
print(f'Date First Seen: {c13_apt1_date}')
print(f'Achilles Heel: {c13_apt1_target}')

# Print Country 13 APT2
print(f'APT Name: {c13_apt2}')
print(f'Goals: {c13_apt2_goal1}, {c13_apt2_goal2}')
print(f'Date First Seen: {c13_apt2_date}')
print(f'Achilles Heel: {c13_apt2_target}')

# Print Country 13 APT3
print(f'APT Name: {c13_apt3}')
print(f'Goals: {c13_apt3_goal1}, {c13_apt3_goal2}')
print(f'Date First Seen: {c13_apt3_date}')
print(f'Achilles Heel: {c13_apt3_target}')

# Print Country 13 APT4
print(f'APT Name: {c13_apt4}')
print(f'Goals: {c13_apt4_goal1}, {c13_apt4_goal2}')
print(f'Date First Seen: {c13_apt4_date}')
print(f'Achilles Heel: {c13_apt4_target}')

# Print Country 13 APT5
print(f'APT Name: {c13_apt5}')
print(f'Goals: {c13_apt5_goal1}, {c13_apt5_goal2}')
print(f'Date First Seen: {c13_apt5_date}')
print(f'Achilles Heel: {c13_apt5_target}')

# Print Country 13 APT6
print(f'APT Name: {c13_apt6}')
print(f'Goals: {c13_apt6_goal1}, {c13_apt6_goal2}')
print(f'Date First Seen: {c13_apt6_date}')
print(f'Achilles Heel: {c13_apt6_target}')

# Print Country 13 APT7
print(f'APT Name: {c13_apt7}')
print(f'Goals: {c13_apt7_goal1}, {c13_apt7_goal2}')
print(f'Date First Seen: {c13_apt7_date}')
print(f'Achilles Heel: {c13_apt7_target}')


# Print Country 14 APT1
print(f'APT Name: {c14_apt1}')
print(f'Goals: {c14_apt1_goal1}, {c14_apt1_goal2}')
print(f'Date First Seen: {c14_apt1_date}')
print(f'Achilles Heel: {c14_apt1_target}')

# Print Country 14 APT2
print(f'APT Name: {c14_apt2}')
print(f'Goals: {c14_apt2_goal1}, {c14_apt2_goal2}')
print(f'Date First Seen: {c14_apt2_date}')
print(f'Achilles Heel: {c14_apt2_target}')

# Print Country 14 APT3
print(f'APT Name: {c14_apt3}')
print(f'Goals: {c14_apt3_goal1}, {c14_apt3_goal2}')
print(f'Date First Seen: {c14_apt3_date}')
print(f'Achilles Heel: {c14_apt3_target}')

# Print Country 14 APT4
print(f'APT Name: {c14_apt4}')
print(f'Goals: {c14_apt4_goal1}, {c14_apt4_goal2}')
print(f'Date First Seen: {c14_apt4_date}')
print(f'Achilles Heel: {c14_apt4_target}')

# Print Country 14 APT5
print(f'APT Name: {c14_apt5}')
print(f'Goals: {c14_apt5_goal1}, {c14_apt5_goal2}')
print(f'Date First Seen: {c14_apt5_date}')
print(f'Achilles Heel: {c14_apt5_target}')

# Print Country 14 APT6
print(f'APT Name: {c14_apt6}')
print(f'Goals: {c14_apt6_goal1}, {c14_apt6_goal2}')
print(f'Date First Seen: {c14_apt6_date}')
print(f'Achilles Heel: {c14_apt6_target}')

# Print Country 14 APT7
print(f'APT Name: {c14_apt7}')
print(f'Goals: {c14_apt7_goal1}, {c14_apt7_goal2}')
print(f'Date First Seen: {c14_apt7_date}')
print(f'Achilles Heel: {c14_apt7_target}')

subprocess.run(f"touch /home/user/challengeServer/dev/*.pdf", shell=True)
subprocess.run(f"touch /home/user/challengeServer/dev/*.docx", shell=True)
subprocess.run(f"tar czf /home/user/challengeServer/dev/apts-and-ips.tar.gz -C /home/user/challengeServer/dev/ ips-*.pdf ips-*.docx", shell=True)

password=subprocess.check_output(f"vmtoolsd --cmd 'info-get guestinfo.password'", shell=True).decode().strip('\n')
command=f"openssl enc -aes-256-ctr -salt -pbkdf2 -in '/home/user/challengeServer/dev/apts-and-ips.tar.gz' -out '/home/user/challengeServer/hosted_files/apts-and-ips.tar.gz.ctr' -k {password}"
subprocess.run(command, shell=True)
subprocess.run(f"rm -rf /home/user/challengeServer/dev/ips-*.pdf /home/user/challengeServer/dev/ips-*.docx", shell=True)
subprocess.run(f"tar czf /home/user/challengeServer/hosted_files/apts.tar.gz -C /home/user/challengeServer/dev/ *.pdf *.docx", shell=True)
 
